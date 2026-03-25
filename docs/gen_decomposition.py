"""Generate the Afarensis Objective-Based System Decomposition document."""
import sys
sys.path.insert(0, "../backend")

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# ── Page Setup ───────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Styles ───────────────────────────────────────────────────────────
style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10)
style.font.color.rgb = RGBColor(0x1F, 0x2A, 0x36)

navy = RGBColor(0x1E, 0x3A, 0x5F)
blue = RGBColor(0x25, 0x63, 0xEB)
dark = RGBColor(0x1F, 0x2A, 0x36)
gray = RGBColor(0x6B, 0x72, 0x80)
white = RGBColor(0xFF, 0xFF, 0xFF)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = navy
        run.font.name = "Calibri"
    return h

def add_para(text, bold=False, italic=False, size=10, color=dark, space_after=6):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.bold = bold
    run.italic = italic
    run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(space_after)
    return p

def add_bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.name = "Calibri"
        run.font.size = Pt(10)
        p.add_run(text).font.name = "Calibri"
    else:
        for run in p.runs:
            run.font.name = "Calibri"
        if not p.runs:
            p.add_run(text).font.name = "Calibri"
    return p

def add_table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(str(text))
        run.font.size = Pt(9)
        run.font.name = "Calibri"
        if header:
            run.bold = True
            run.font.color.rgb = white
            shading = cell._element.get_or_add_tcPr()
            bg = shading.makeelement(qn("w:shd"), {
                qn("w:val"): "clear",
                qn("w:color"): "auto",
                qn("w:fill"): "1E3A5F"
            })
            shading.append(bg)
        else:
            run.font.color.rgb = dark

def make_table(headers, rows):
    t = doc.add_table(rows=0, cols=len(headers))
    t.style = "Table Grid"
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(t, headers, header=True)
    for row in rows:
        add_table_row(t, row)
    return t

# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("AFARENSIS")
run.font.size = Pt(36)
run.font.color.rgb = navy
run.bold = True
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Objective-Based System Decomposition")
run.font.size = Pt(20)
run.font.color.rgb = blue
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Emergent Flows Architecture")
run.font.size = Pt(14)
run.font.color.rgb = gray
run.italic = True
run.font.name = "Calibri"

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Synthetic Ascension")
run.font.size = Pt(11)
run.font.color.rgb = dark
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f"Version 2.1  |  {datetime.now().strftime('%B %Y')}  |  CONFIDENTIAL")
run.font.size = Pt(10)
run.font.color.rgb = gray
run.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# DOCUMENT CONTROL
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Document Control", 1)

make_table(
    ["Field", "Value"],
    [
        ["Document Title", "Afarensis Objective-Based System Decomposition"],
        ["Version", "2.1"],
        ["Classification", "Confidential"],
        ["Author", "Synthetic Ascension Systems Architecture"],
        ["Date", datetime.now().strftime("%Y-%m-%d")],
        ["Status", "Production Release"],
        ["Applicable Standard", "ICH E9(R1), 21 CFR Part 11, CDISC ADaM/SDTM"],
    ],
)

doc.add_paragraph()
add_para(
    "This document defines every feature, capability, and experience of the Afarensis platform "
    "strictly from first principles. Flows are not declared; they emerge naturally from objectives, "
    "features, dependencies, and constraints.",
    italic=True, color=gray, size=10
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CORE PRINCIPLE
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Core Decomposition Principle", 1)

add_para(
    "Every part of the system is defined as:", bold=True, size=11
)
add_para(
    "Objective \u2192 Feature \u2192 Mechanism \u2192 Inputs \u2192 Outputs \u2192 "
    "Constraints \u2192 Failure Modes \u2192 Observability",
    bold=True, size=11, color=blue
)
add_para(
    "Flows are NOT defined directly. They are inferred artifacts of the system's structure. "
    "A reader should be able to derive all possible system flows, identify bottlenecks, "
    "detect circular dependencies, and understand execution order purely from this document."
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 1: SYSTEM OBJECTIVE MAP
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Section 1: System Objective Map", 1)

add_para(
    "Each objective represents a real, independently meaningful unit of work. "
    "Objectives are falsifiable: either the system fulfills them or it does not.",
    italic=True, color=gray
)

objectives = [
    ("OBJ-01", "Study Protocol Ingestion", "Accept a clinical study protocol (PDF/DOCX/text), parse it into structured fields (indication, population, endpoints, covariates, inclusion/exclusion criteria), and persist the specification for downstream use.", "Critical"),
    ("OBJ-02", "External Evidence Discovery", "Query PubMed, ClinicalTrials.gov, and Semantic Scholar APIs using structured search terms derived from the parsed protocol, retrieve matching publications, and store them as evidence records with provenance metadata.", "Critical"),
    ("OBJ-03", "Evidence Comparability Scoring", "For each evidence record, compute a multi-dimensional comparability score (population similarity, endpoint alignment, covariate coverage, temporal alignment, evidence quality) against the reference study specification.", "Critical"),
    ("OBJ-04", "Bias Detection and Quantification", "Identify and quantify five bias types (selection, confounding, measurement, temporal, publication) for each evidence-comparability pair. Compute fragility indices and E-values for unmeasured confounding sensitivity.", "Critical"),
    ("OBJ-05", "Causal Inference Estimation", "Execute causal inference methods (Cox PH, IPTW, propensity scoring, Kaplan-Meier) on study data to produce treatment effect estimates with confidence intervals and p-values.", "Critical"),
    ("OBJ-06", "Missing Data Handling", "Implement multiple imputation (MICE), MMRM, and tipping-point analyses to assess the impact of missing data on primary and sensitivity analyses.", "Critical"),
    ("OBJ-07", "Multiplicity Adjustment", "Apply family-wise error rate control (Bonferroni, Holm, Hochberg, Benjamini-Hochberg) across multiple hypothesis tests with hierarchical gatekeeping.", "Supporting"),
    ("OBJ-08", "Bayesian Inference", "Perform Bayesian analysis with prior elicitation, MCMC posterior sampling, and credible interval computation for regulatory submissions requiring Bayesian evidence.", "Supporting"),
    ("OBJ-09", "Interim Analysis", "Calculate group-sequential boundaries (O'Brien-Fleming, Pocock, Lan-DeMets), evaluate conditional power at interim looks, and generate DSMB reporting packages.", "Supporting"),
    ("OBJ-10", "CDISC ADaM Dataset Generation", "Generate CDISC ADaM-compliant analysis datasets (ADSL, ADAE, ADTTE, ADLB) from study data, with variable-level metadata, codelists, and validation against ADaM Implementation Guide rules.", "Critical"),
    ("OBJ-11", "CDISC SDTM Dataset Generation", "Generate CDISC SDTM-compliant tabulation datasets (DM, AE, LB, VS, EX, DS) and annotated CRFs from raw study data.", "Critical"),
    ("OBJ-12", "Tables, Figures, and Listings Generation", "Produce publication-quality TFLs: demographics tables, adverse event summaries, Kaplan-Meier curves, forest plots, and covariate balance Love plots.", "Critical"),
    ("OBJ-13", "Statistical Analysis Plan Authoring", "Generate a formal SAP document following ICH E9(R1) structure with estimand framework, multiplicity strategy, missing data plan, and TFL shell specifications.", "Critical"),
    ("OBJ-14", "Clinical Study Report Generation", "Produce ICH E3-compliant CSR sections: Synopsis, Section 11 (efficacy), Section 12 (safety), Appendix 16 (statistical methods).", "Critical"),
    ("OBJ-15", "Analysis Data Reviewer Guide Generation", "Generate a PhUSE-compliant ADRG documenting the computational environment, dataset structure, programs, and variable definitions for FDA reviewer reproducibility.", "Critical"),
    ("OBJ-16", "Define-XML Generation", "Produce CDISC Define-XML 2.1 metadata documents describing all ADaM datasets, variables, origins, and codelists.", "Critical"),
    ("OBJ-17", "eCTD Module 5 Packaging", "Assemble all analysis outputs into FDA eCTD Module 5 directory structure with Study Tagging Files, document checksums, and manifest.", "Critical"),
    ("OBJ-18", "Collaborative Evidence Review", "Enable multi-reviewer workflows with assignment, commenting, real-time presence, conflict resolution, and cryptographic e-signature on decisions.", "Critical"),
    ("OBJ-19", "Audit Trail and Regulatory Traceability", "Record every system action (creation, modification, deletion, decision, generation) with user attribution, timestamp, IP address, old/new values, and 7-year retention.", "Critical"),
    ("OBJ-20", "Multi-Tenant Organization Isolation", "Ensure complete data isolation between organizations: users, projects, evidence, and artifacts are scoped to the owning organization with no cross-org data leakage.", "Critical"),
    ("OBJ-21", "Identity and Access Management", "Authenticate users via JWT with refresh token rotation, enforce role-based access control (Admin/Reviewer/Analyst/Viewer), and provide password reset via verified email codes.", "Critical"),
    ("OBJ-22", "Submission Readiness Assessment", "Compute a regulatory readiness score across all project dimensions (protocol lock, evidence sufficiency, bias assessment, TFL completeness, CSR sections, eCTD packaging) and surface gaps.", "Supporting"),
    ("OBJ-23", "Reproducibility Verification", "Track computational environment (software versions, package manifests, random seeds) and enable independent reconstruction of all analysis outputs.", "Critical"),
    ("OBJ-24", "System Observability", "Provide request-level metrics (latency percentiles, error rates, throughput), Sentry error tracking, background task monitoring, cache hit ratios, and database health.", "Supporting"),
]

make_table(
    ["ID", "Objective", "Description", "Criticality"],
    [[o[0], o[1], o[2][:120] + "..." if len(o[2]) > 120 else o[2], o[3]] for o in objectives],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 2: FEATURE DECOMPOSITION (selected critical objectives)
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Section 2: Feature Decomposition", 1)

add_para(
    "Each feature is decomposed into: Objective Alignment, Mechanism, Inputs, Outputs, "
    "User Interaction Surface, Constraints, Failure Modes, and Observability.",
    italic=True, color=gray
)

# Feature: OBJ-05 Causal Inference
add_heading_styled("OBJ-05: Causal Inference Estimation", 2)

add_heading_styled("Feature 5.1: Cox Proportional Hazards Model", 3)

add_para("Objective Alignment: ", bold=True, size=10)
add_bullet("Supports OBJ-05 (Critical). Primary analysis method for time-to-event endpoints.")

add_para("Mechanism (step-by-step):", bold=True, size=10)
add_bullet("1. Receive input arrays: time_to_event (float[]), event_indicator (bool[]), treatment (binary[]), covariates (float[][]).")
add_bullet("2. Construct partial likelihood function L(\u03b2) = \u220f [exp(\u03b2\u1d40x_i) / \u2211_{j\u2208R(t_i)} exp(\u03b2\u1d40x_j)] for each event time t_i.")
add_bullet("3. Compute gradient (score function) and Hessian (observed information matrix) analytically.")
add_bullet("4. Iterate Newton-Raphson: \u03b2_{k+1} = \u03b2_k - H\u207b\u00b9(\u03b2_k) \u00b7 U(\u03b2_k) until ||\u0394\u03b2|| < 1e-8 or max 100 iterations.")
add_bullet("5. Extract hazard ratios: HR = exp(\u03b2), 95% CI from Wald: exp(\u03b2 \u00b1 1.96 \u00b7 SE(\u03b2)).")
add_bullet("6. Compute concordance index (Harrell\u2019s C) by comparing predicted risk rankings to observed event order.")
add_bullet("7. Evaluate proportional hazards assumption via Schoenfeld residual correlation with time.")

add_para("Inputs:", bold=True, size=10)
make_table(
    ["Input", "Type", "Origin", "Validation"],
    [
        ["time_to_event", "float[]", "ADaM ADTTE.AVAL", "> 0, no NaN"],
        ["event_indicator", "int[] (0/1)", "ADaM ADTTE.CNSR (inverted)", "Binary only"],
        ["treatment", "int[] (0/1)", "ADaM ADSL.TRT01P", "Binary only"],
        ["covariates", "float[][]", "ADaM ADSL (AGE, SEX, etc.)", "No constant columns"],
    ],
)

add_para("Outputs:", bold=True, size=10)
make_table(
    ["Output", "Type", "Downstream Consumer"],
    [
        ["coefficients", "dict {var: beta}", "Forest plot (OBJ-12)"],
        ["hazard_ratios", "dict {var: HR}", "CSR Section 11 (OBJ-14)"],
        ["confidence_intervals", "dict {var: (lo, hi)}", "Forest plot, SAP (OBJ-13)"],
        ["p_values", "dict {var: float}", "Multiplicity adjustment (OBJ-07)"],
        ["concordance_index", "float (0-1)", "Model diagnostics, ADRG (OBJ-15)"],
        ["schoenfeld_p", "float", "PH assumption check, audit (OBJ-19)"],
    ],
)

add_para("Constraints:", bold=True, size=10)
add_bullet("Statistical: ", bold_prefix="Statistical: ")
add_bullet("Requires \u226510 events per covariate (rule of thumb) for stable estimation.")
add_bullet("Proportional hazards assumption must hold (Schoenfeld p > 0.05).")
add_bullet("Data availability: ", bold_prefix="Data: ")
add_bullet("Requires complete time-to-event and censoring data; missing values handled by OBJ-06.")

add_para("Failure Modes:", bold=True, size=10)
make_table(
    ["Condition", "Type", "Impact", "Mitigation"],
    [
        ["Newton-Raphson non-convergence", "Surfaced (error)", "No HR estimate produced", "Fall back to Firth correction or report failure"],
        ["Singular Hessian", "Surfaced (warning)", "SE/CI unavailable", "Profile likelihood CI or bootstrap"],
        ["Zero events in one arm", "Surfaced (error)", "HR undefined", "Report descriptive statistics only"],
        ["PH assumption violated", "Surfaced (warning)", "HR interpretation invalid", "Switch to time-varying or restricted mean survival"],
    ],
)

add_para("Observability:", bold=True, size=10)
add_bullet("Logged: convergence iterations, final log-likelihood, concordance, PH test result.")
add_bullet("Audit: full coefficient vector, input data hash (SHA-256), software version, random seed.")
add_bullet("Reconstructible: any reviewer can re-execute with stored inputs + parameters.")

doc.add_page_break()

# Feature: OBJ-10 ADaM
add_heading_styled("OBJ-10: CDISC ADaM Dataset Generation", 2)
add_heading_styled("Feature 10.1: ADSL (Subject-Level Analysis Dataset)", 3)

add_para("Mechanism:", bold=True, size=10)
add_bullet("1. Query project cohort criteria and study definition from processing_config.")
add_bullet("2. Generate subject-level records with CDISC-required variables (STUDYID, USUBJID, SUBJID, SITEID, ARM, TRT01P, TRT01A).")
add_bullet("3. Derive population flags (ITTFL, SAFFL, COMPLFL) from inclusion/exclusion criteria.")
add_bullet("4. Compute derived variables (AGE, AGEGR1, TRTDUR, DTHFL) from raw demographics + dates.")
add_bullet("5. Validate all variables against ADaM IG: name \u22648 chars, label \u226440 chars, required variables present.")
add_bullet("6. Store dataset + variable metadata in adam_datasets table.")

add_para("Outputs:", bold=True, size=10)
make_table(
    ["Output", "Format", "Consumer"],
    [
        ["ADSL dataset", "JSON (rows + variable specs)", "TFL generation (OBJ-12), Cox PH (OBJ-05)"],
        ["Variable metadata", "JSON (name, label, type, origin, codelist)", "Define-XML (OBJ-16)"],
        ["Validation report", "JSON (pass/fail per rule)", "ADRG (OBJ-15), audit (OBJ-19)"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 3: DEPENDENCY GRAPH
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Section 3: Dependency Graph", 1)

add_para(
    "This directed graph defines what each objective depends on and what depends on it. "
    "From this graph, all possible system flows can be derived.",
    italic=True, color=gray
)

deps = [
    ("OBJ-01", "Protocol Ingestion", "None (entry point)", "OBJ-02, OBJ-03, OBJ-10, OBJ-11, OBJ-13"),
    ("OBJ-02", "Evidence Discovery", "OBJ-01", "OBJ-03, OBJ-04, OBJ-18"),
    ("OBJ-03", "Comparability Scoring", "OBJ-01, OBJ-02", "OBJ-04, OBJ-05, OBJ-18"),
    ("OBJ-04", "Bias Detection", "OBJ-03", "OBJ-05, OBJ-06, OBJ-22"),
    ("OBJ-05", "Causal Inference", "OBJ-04, OBJ-10", "OBJ-07, OBJ-12, OBJ-14"),
    ("OBJ-06", "Missing Data", "OBJ-05, OBJ-10", "OBJ-05 (cycle: sensitivity), OBJ-14"),
    ("OBJ-07", "Multiplicity", "OBJ-05", "OBJ-12, OBJ-13, OBJ-14"),
    ("OBJ-08", "Bayesian", "OBJ-05, OBJ-10", "OBJ-14, OBJ-09"),
    ("OBJ-09", "Interim Analysis", "OBJ-05, OBJ-08", "OBJ-14"),
    ("OBJ-10", "ADaM Generation", "OBJ-01", "OBJ-05, OBJ-11, OBJ-12, OBJ-16"),
    ("OBJ-11", "SDTM Generation", "OBJ-01, OBJ-10", "OBJ-16, OBJ-17"),
    ("OBJ-12", "TFL Generation", "OBJ-05, OBJ-07, OBJ-10", "OBJ-14, OBJ-17"),
    ("OBJ-13", "SAP Authoring", "OBJ-01, OBJ-07", "OBJ-14, OBJ-17, OBJ-22"),
    ("OBJ-14", "CSR Generation", "OBJ-05, OBJ-07, OBJ-12", "OBJ-17"),
    ("OBJ-15", "ADRG Generation", "OBJ-10, OBJ-23", "OBJ-17"),
    ("OBJ-16", "Define-XML", "OBJ-10, OBJ-11", "OBJ-17"),
    ("OBJ-17", "eCTD Packaging", "OBJ-14, OBJ-15, OBJ-16", "OBJ-22 (terminal)"),
    ("OBJ-18", "Collaborative Review", "OBJ-02, OBJ-03", "OBJ-19, OBJ-22"),
    ("OBJ-19", "Audit Trail", "All objectives (passive)", "OBJ-22, OBJ-23"),
    ("OBJ-20", "Tenant Isolation", "OBJ-21 (identity)", "All data access paths"),
    ("OBJ-21", "IAM", "None (entry point)", "OBJ-20, all authenticated endpoints"),
    ("OBJ-22", "Readiness Assessment", "OBJ-04, OBJ-13, OBJ-17, OBJ-18", "None (terminal)"),
    ("OBJ-23", "Reproducibility", "OBJ-19, OBJ-10", "OBJ-15"),
    ("OBJ-24", "Observability", "None (passive)", "None (operational)"),
]

make_table(
    ["Objective", "Name", "Upstream Dependencies", "Downstream Consumers"],
    [[d[0], d[1], d[2], d[3]] for d in deps],
)

doc.add_paragraph()
add_para("Critical Path: OBJ-01 \u2192 OBJ-02 \u2192 OBJ-03 \u2192 OBJ-04 \u2192 OBJ-05 \u2192 OBJ-07 \u2192 OBJ-12 \u2192 OBJ-14 \u2192 OBJ-17", bold=True, color=blue)
add_para("Parallel Path: OBJ-01 \u2192 OBJ-10 \u2192 OBJ-16 \u2192 OBJ-17", bold=True, color=blue)
add_para("Cycle: OBJ-05 \u2194 OBJ-06 (sensitivity analyses feed back into primary analysis)", italic=True, color=gray)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 4: EMERGENT EXPERIENCE SURFACES
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Section 4: Emergent Experience Surfaces", 1)

add_para(
    "These interaction patterns are NOT prescribed flows. They emerge from the dependency graph.",
    italic=True, color=gray
)

surfaces = [
    ("Study Architect", "Biostatistician defines protocol, covariates, cohort criteria, and analysis spec.",
     "Protocol text upload, covariate selection, endpoint definition, estimand specification.",
     "OBJ-01 \u2192 processing_config populated \u2192 downstream objectives unblocked."),
    ("Evidence Curator", "Analyst discovers, scores, and critiques external evidence.",
     "Search query refinement, comparability threshold selection, evidence acceptance/rejection.",
     "OBJ-02 \u2192 OBJ-03 \u2192 OBJ-04 chain. Human decides which evidence to include."),
    ("Statistical Analyst", "Analyst configures and runs statistical models, reviews diagnostics.",
     "Model selection, covariate set, sensitivity parameters, missing data strategy.",
     "OBJ-05/06/07/08/09 constellation. Human selects methods, system computes."),
    ("Regulatory Reviewer", "Reviewer evaluates evidence quality, submits e-signed decisions.",
     "Accept/reject/defer decision, confidence level, written rationale.",
     "OBJ-18 \u2192 OBJ-19. Human-in-the-loop is unavoidable for regulatory decisions."),
    ("Document Assembler", "Analyst generates TFLs, CSR sections, and packages for submission.",
     "TFL selection, CSR section ordering, eCTD structure confirmation.",
     "OBJ-12 \u2192 OBJ-14 \u2192 OBJ-17. Mostly automated, human reviews output."),
    ("Auditor", "Regulator or QA traces any output back to its inputs and transformations.",
     "Audit log filtering, artifact download, reproducibility verification.",
     "OBJ-19 \u2192 OBJ-23. Read-only, no human decisions required."),
]

for name, desc, decisions, emergence in surfaces:
    add_heading_styled(f"Surface: {name}", 3)
    add_para(desc)
    add_para("Decision points: ", bold=True, size=10)
    add_para(decisions, italic=True)
    add_para("Emergence: ", bold=True, size=10)
    add_para(emergence, color=blue)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SECTION 5: EDGE CASE & SYSTEM STRESS BEHAVIOR
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Section 5: System Stress Map", 1)

stress = [
    ("Missing protocol text", "OBJ-01", "Parser returns empty specification", "User prompted to upload document or enter manually", "OBJ-02, OBJ-10 blocked"),
    ("Zero PubMed results", "OBJ-02", "No evidence records created", "System reports no matches; user can widen search terms", "OBJ-03 has nothing to score"),
    ("All evidence scores < 0.3", "OBJ-03", "No viable comparators", "Warning surfaced; regulatory readiness drops to 0", "OBJ-05 proceeds but with caveats"),
    ("Cox model non-convergence", "OBJ-05", "No HR estimate", "Error logged, fallback to descriptive statistics", "OBJ-07, OBJ-12 receive null estimates"),
    ("100% missing data in a covariate", "OBJ-06", "Imputation impossible for that variable", "Variable excluded from model; audit log records decision", "Reduced covariate set in OBJ-05"),
    ("All p-values > 0.05 after multiplicity", "OBJ-07", "No significant findings", "Results reported as-is; no artificial significance", "CSR reports non-significance"),
    ("PubMed API rate limit exceeded", "OBJ-02", "429 response from NCBI", "Exponential backoff with retry (3 attempts)", "Task remains in RUNNING state"),
    ("S3 storage unavailable", "OBJ-17", "Artifact not persisted to cloud", "Local copy preserved; warning logged", "Download still works from local"),
    ("Reviewer conflict (split decision)", "OBJ-18", "No consensus on evidence", "Escalation to senior reviewer; conflict resolution UI", "Decision remains PENDING"),
    ("Database connection pool exhausted", "OBJ-24", "New requests queued/rejected", "Pool pre-ping detects; health endpoint reports degraded", "Rate limiting kicks in"),
]

make_table(
    ["Stress Condition", "Objective", "System Behavior", "User Impact", "Downstream Effect"],
    stress,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# ADDENDUM A: OPERATIONAL MODEL
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Addendum A: Operational Readiness Model", 1)

add_para(
    "For the system to function end-to-end, the following conditions must be true:",
    bold=True
)

readiness = [
    ("Infrastructure", "PostgreSQL database accessible, Alembic migrations applied, SECRET_KEY configured (non-default)"),
    ("Authentication", "At least one ADMIN user exists, JWT signing key is unique per environment"),
    ("Organization", "At least one Organization record exists, users are assigned to it"),
    ("Protocol Input", "At least one Project created with research_intent populated"),
    ("Evidence Pipeline", "PubMed/ClinicalTrials API accessible OR uploaded documents provided"),
    ("Statistical Engine", "NumPy, SciPy, statsmodels importable; simulation data available for empty projects"),
    ("Document Engine", "python-docx importable; artifact directory writable; storage backend (local or S3) configured"),
    ("Audit System", "Audit logging enabled (ENABLE_AUDIT_LOG=true), 7-year retention policy active"),
    ("Observability", "Health endpoint responsive (/api/v1/health), metrics middleware active"),
]

make_table(
    ["Domain", "Readiness Condition"],
    readiness,
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# ADDENDUM B: RECONSTRUCTION MODEL
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("Addendum B: Reconstruction Model", 1)

add_para(
    "A regulator or auditor can reconstruct any output through the following chain:",
    bold=True
)

add_heading_styled("Output Reconstruction", 3)
add_bullet("Every regulatory artifact (CSR, SAR, TFL) stores its generation timestamp, input data hash (SHA-256), software version, and generating user ID.")
add_bullet("The audit log captures every mutation (create, update, delete) with old_values and new_values JSON fields.")
add_bullet("Statistical analysis results include: model specification, convergence diagnostics, random seed, full coefficient vector, and input data fingerprint.")
add_bullet("ADaM/SDTM datasets include variable-level metadata (origin, derivation rule, codelist reference) traceable to Define-XML.")
add_bullet("The reproducibility manifest records: Python version, package versions (pinned), operating system, and execution timestamp.")

add_heading_styled("Decision Tracing", 3)
add_bullet("Every review decision records: reviewer ID, decision value, confidence level, written rationale, cryptographic signature (SHA-256 of decision + timestamp + user ID), and review duration.")
add_bullet("Conflict resolution records the original split decisions, the resolution rationale, and the resolving authority.")
add_bullet("E-signature chain: decision hash \u2192 audit log entry \u2192 regulatory artifact reference.")

add_heading_styled("Transformation Validation", 3)
add_bullet("Raw input \u2192 Parsed specification: diffable (old_values vs new_values in audit log).")
add_bullet("Evidence search \u2192 Results: query string, API response count, and retrieval ranks stored per record.")
add_bullet("Statistical model \u2192 Results: full reproducibility via stored inputs + parameters + random seed.")
add_bullet("Analysis results \u2192 TFL: direct mapping from result dict keys to table cells/figure data points.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# SYSTEM INVENTORY SUMMARY
# ══════════════════════════════════════════════════════════════════════
add_heading_styled("System Inventory Summary", 1)

make_table(
    ["Dimension", "Count"],
    [
        ["System Objectives", "24"],
        ["Database Models", "24 (with 5 enums)"],
        ["API Endpoints", "142 RESTful + 1 WebSocket"],
        ["Service Classes", "30+"],
        ["Statistical Methods", "20+ (Cox PH, IPTW, KM, Bayesian, MMRM, meta-analysis, etc.)"],
        ["Document Types", "10+ (CSR, ADRG, eCTD, Define-XML, SAR, SAP, TFL)"],
        ["CDISC Domains", "10 (ADaM: ADSL, ADAE, ADTTE, ADLB; SDTM: DM, AE, LB, VS, EX, DS)"],
        ["Frontend Pages", "32 React/TypeScript"],
        ["Configuration Options", "170+"],
        ["Test Cases", "316 passing"],
        ["External Integrations", "6 (PubMed, ClinicalTrials, Semantic Scholar, Claude, GPT-4, Gemini)"],
    ],
)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("A blueprint where flows are shadows cast by structure.")
run.italic = True
run.font.size = Pt(11)
run.font.color.rgb = navy
run.font.name = "Calibri"

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Afarensis Enterprise v2.1 \u2014 Synthetic Ascension \u2014 Confidential")
run.font.size = Pt(9)
run.font.color.rgb = gray
run.font.name = "Calibri"

# ── Save ─────────────────────────────────────────────────────────────
output = "Afarensis_System_Decomposition.docx"
doc.save(output)
print(f"Generated: {output}")
