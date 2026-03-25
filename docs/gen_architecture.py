"""Generate Architecture, Data Flow, and ERD diagrams document."""
import sys
sys.path.insert(0, "../backend")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()
for s in doc.sections:
    s.page_width = Inches(8.5); s.page_height = Inches(11)
    s.top_margin = Inches(0.9); s.bottom_margin = Inches(0.7)
    s.left_margin = Inches(0.8); s.right_margin = Inches(0.8)

navy = RGBColor(0x1E, 0x3A, 0x5F)
blue = RGBColor(0x25, 0x63, 0xEB)
dark = RGBColor(0x1F, 0x2A, 0x36)
gray = RGBColor(0x6B, 0x72, 0x80)
white = RGBColor(0xFF, 0xFF, 0xFF)
green = RGBColor(0x05, 0x96, 0x69)

style = doc.styles["Normal"]
style.font.name = "Calibri"; style.font.size = Pt(10); style.font.color.rgb = dark

def h1(t):
    h = doc.add_heading(t, 1)
    for r in h.runs: r.font.color.rgb = navy; r.font.name = "Calibri"
def h2(t):
    h = doc.add_heading(t, 2)
    for r in h.runs: r.font.color.rgb = navy; r.font.name = "Calibri"
def h3(t):
    h = doc.add_heading(t, 3)
    for r in h.runs: r.font.color.rgb = blue; r.font.name = "Calibri"
def p(t, **kw):
    para = doc.add_paragraph()
    r = para.add_run(t)
    r.font.size = Pt(kw.get("size", 10)); r.font.color.rgb = kw.get("color", dark)
    r.bold = kw.get("bold", False); r.italic = kw.get("italic", False); r.font.name = "Calibri"
    para.paragraph_format.space_after = Pt(kw.get("after", 6))
    return para
def bullet(t, bp=None):
    para = doc.add_paragraph(style="List Bullet")
    if bp:
        r = para.add_run(bp); r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10)
        para.add_run(t).font.name = "Calibri"
    else:
        para.add_run(t).font.name = "Calibri"
def code(t):
    para = doc.add_paragraph()
    r = para.add_run(t); r.font.name = "Consolas"; r.font.size = Pt(8); r.font.color.rgb = dark
    para.paragraph_format.space_after = Pt(4)
def add_row(tbl, cells, hdr=False):
    row = tbl.add_row()
    for i, t in enumerate(cells):
        c = row.cells[i]; c.text = ""; pp = c.paragraphs[0]
        r = pp.add_run(str(t)); r.font.size = Pt(8); r.font.name = "Calibri"
        if hdr:
            r.bold = True; r.font.color.rgb = white
            sh = c._element.get_or_add_tcPr()
            sh.append(sh.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "1E3A5F"}))
        else:
            r.font.color.rgb = dark
def tbl(hdrs, rows):
    t = doc.add_table(rows=0, cols=len(hdrs)); t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_row(t, hdrs, True)
    for row in rows: add_row(t, row)
    doc.add_paragraph(); return t

# ══════════════════════════════════════════════════════════════════
# TITLE
# ══════════════════════════════════════════════════════════════════
for _ in range(5): doc.add_paragraph()
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("AFARENSIS ENTERPRISE"); r.font.size = Pt(32); r.font.color.rgb = navy; r.bold = True; r.font.name = "Calibri"
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("Architecture, Data Flow & ERD"); r.font.size = Pt(18); r.font.color.rgb = blue; r.font.name = "Calibri"
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run(f"Version 2.1  |  {datetime.now().strftime('%B %Y')}  |  Confidential"); r.font.size = Pt(10); r.font.color.rgb = gray; r.font.name = "Calibri"
doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# C1: SYSTEM CONTEXT
# ══════════════════════════════════════════════════════════════════
h1("C1: System Context Diagram")
p("The C1 diagram shows how Afarensis Enterprise interacts with external actors and systems.", italic=True, color=gray)

doc.add_paragraph()
code("""
+-----------------------------------------------------------------------------------+
|                              EXTERNAL ACTORS                                       |
|                                                                                   |
|   [Biostatistician]    [Reviewer]    [Admin]    [FDA/EMA Regulator]               |
|         |                  |            |              |                           |
+---------|------------------|------------|--------------|---------------------------+
          |                  |            |              |
          v                  v            v              v
+-----------------------------------------------------------------------------------+
|                                                                                   |
|                        AFARENSIS ENTERPRISE v2.1                                  |
|                   Regulatory Evidence Review Platform                             |
|                                                                                   |
|   - Study protocol management                                                    |
|   - External evidence discovery & scoring                                         |
|   - Statistical analysis (Cox PH, IPTW, KM, Bayesian)                            |
|   - CDISC dataset generation (ADaM, SDTM)                                        |
|   - Regulatory document generation (CSR, ADRG, SAP, TFL)                         |
|   - eCTD Module 5 submission packaging                                            |
|   - Collaborative multi-reviewer workflows                                        |
|   - 21 CFR Part 11 audit trail                                                    |
|                                                                                   |
+---------|-----------|-----------|-----------|-----------|-------------------------+
          |           |           |           |           |
          v           v           v           v           v
+----------+ +----------+ +-----------+ +----------+ +---------+
|  PubMed  | |ClinTrials| | Semantic  | | Claude / | | SendGrid|
|  (NCBI)  | |  .gov    | | Scholar   | | GPT-4    | | (SMTP)  |
+----------+ +----------+ +-----------+ +----------+ +---------+
""")

h2("External Actors")
tbl(["Actor", "Role", "Interactions"],
    [["Biostatistician", "Primary user", "Defines protocol, configures analysis, generates datasets & documents"],
     ["Reviewer", "Evidence evaluator", "Reviews evidence, submits e-signed decisions, resolves conflicts"],
     ["Admin", "System administrator", "Manages users/orgs, monitors system health, configures settings"],
     ["Regulator (FDA/EMA)", "Downstream consumer", "Receives eCTD packages, reviews CSR/ADRG/Define-XML"]])

h2("External Systems")
tbl(["System", "Protocol", "Purpose", "Rate Limit"],
    [["PubMed (NCBI)", "REST/XML", "Literature search & retrieval", "3 req/sec"],
     ["ClinicalTrials.gov", "REST/JSON", "Clinical trial data", "2 req/sec"],
     ["Semantic Scholar", "REST/JSON", "Academic paper search & recommendations", "10 req/sec"],
     ["Claude / GPT-4", "REST/JSON", "Evidence critique, entity extraction, summarization", "60 req/min"],
     ["SendGrid / SMTP", "SMTP/TLS", "Password reset emails, notifications", "On-demand"]])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# C2: CONTAINER DIAGRAM
# ══════════════════════════════════════════════════════════════════
h1("C2: Container Diagram")
p("The C2 diagram shows the major containers (deployable units) and their interactions.", italic=True, color=gray)

doc.add_paragraph()
code("""
+------------------------------------------------------------------+
|                        BROWSER (Client)                           |
|                                                                   |
|  +------------------------------------------------------------+  |
|  |              React SPA (TypeScript + Vite)                  |  |
|  |  - 22 workflow pages (10-step regulatory pipeline)          |  |
|  |  - Tailwind CSS + lucide-react icons                        |  |
|  |  - JWT stored in memory (not localStorage)                  |  |
|  |  - WebSocket for real-time collaboration                    |  |
|  +------------------------------|-----------------------------+  |
+----------------------------------|--------------------------------+
                                   | HTTPS (JSON + WebSocket)
                                   v
+------------------------------------------------------------------+
|                   FASTAPI BACKEND (Python 3.10)                   |
|                                                                   |
|  +------------------+  +------------------+  +-----------------+  |
|  | API Layer        |  | Service Layer    |  | Core Layer      |  |
|  | (142 REST + 1 WS)|  | (24 services)    |  | (10 modules)    |  |
|  |                  |  |                  |  |                 |  |
|  | Auth (8)         |  | StatisticalModels|  | Security (JWT)  |  |
|  | Projects (6)     |  | DocumentGenerator|  | Database (async)|  |
|  | Evidence (6)     |  | AdamService      |  | Cache (Redis)   |  |
|  | Review (9)       |  | SDTMService      |  | RateLimiter     |  |
|  | Artifacts (14)   |  | TFLGenerator     |  | Storage (S3)    |  |
|  | Study (27)       |  | CSRGenerator     |  | Observability   |  |
|  | Search (10)      |  | BayesianMethods  |  | Pagination      |  |
|  | CDISC (10)       |  | InterimAnalysis  |  | TaskQueue       |  |
|  | Statistics (6)   |  | EmailService     |  | Exceptions      |  |
|  | Admin (9)        |  | ExternalAPIs     |  | Config          |  |
|  | System (7)       |  | TaskQueue        |  |                 |  |
|  | Tasks (4)        |  | CollabReview     |  |                 |  |
|  +------------------+  +------------------+  +-----------------+  |
|                                                                   |
|  Middleware Stack: SecurityHeaders > Metrics > CORS > TrustedHost |
+-----------|---------------|---------------|-----------|------------+
            |               |               |           |
            v               v               v           v
    +------------+   +------------+   +-----------+  +-----------+
    | PostgreSQL |   |   Redis    |   |  S3 / FS  |  | Sentry    |
    | 16-alpine  |   | 7-alpine   |   | Storage   |  | (errors)  |
    | + PgBouncer|   |            |   |           |  |           |
    +------------+   +------------+   +-----------+  +-----------+
""")

h2("Container Details")
tbl(["Container", "Technology", "Purpose", "Port"],
    [["React SPA", "React 18 + TypeScript + Vite", "Single-page application served as static files from FastAPI", "8000 (same origin)"],
     ["FastAPI Backend", "Python 3.10 + uvicorn", "API server, business logic, document generation", "8000"],
     ["PostgreSQL", "PostgreSQL 16 Alpine", "Primary data store (24 tables, Alembic migrations)", "5432"],
     ["PgBouncer", "edoburu/pgbouncer", "Connection pooling (transaction mode, 200 max clients)", "6432"],
     ["Redis", "Redis 7 Alpine", "Caching, rate limiting, session store", "6379"],
     ["S3 / Local FS", "AWS S3 / MinIO / local", "Artifact storage (PDFs, DOCX, HTML, datasets)", "N/A"],
     ["Sentry", "sentry.io (SaaS)", "Error tracking, APM, performance monitoring", "N/A"],
     ["pg-backup", "postgres-backup-local", "Automated daily PostgreSQL backups (30-day retention)", "N/A"]])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# C3: COMPONENT DIAGRAM
# ══════════════════════════════════════════════════════════════════
h1("C3: Component Diagram (Backend)")
p("The C3 diagram decomposes the FastAPI backend into its internal components.", italic=True, color=gray)

h2("API Layer (142 endpoints)")
tbl(["Group", "Endpoints", "Auth Required", "Rate Limited"],
    [["Authentication", "8 (login, logout, refresh, password reset)", "Partial (login/reset public)", "Yes (5/min login, 3/5min reset)"],
     ["Projects", "6 (CRUD, upload, discover)", "Yes", "Yes (20/min create, 5/min discover)"],
     ["Evidence & Analysis", "6 (anchors, comparability, bias, critique)", "Yes", "Yes (5/min bias)"],
     ["Review & Decisions", "9 (assignments, comments, decisions, conflicts)", "Yes", "No"],
     ["Study Configuration", "27 (GET/PUT for 5 sections + compute endpoints)", "Yes", "No"],
     ["Regulatory Artifacts", "14 (CSR, ADRG, Define-XML, eCTD, SAP)", "Yes", "Yes (10/min generate)"],
     ["CDISC Datasets", "10 (ADaM + SDTM generate/validate)", "Yes", "No"],
     ["Search", "10 (semantic, hybrid, citation, Semantic Scholar)", "Yes", "Yes (30/min search)"],
     ["Statistics", "6 (full analysis, Bayesian, interim)", "Yes", "No"],
     ["Background Tasks", "4 (status, result, list, cancel)", "Yes", "No"],
     ["System / Admin", "7 (metrics, health, cache, storage, users)", "Admin only", "No"],
     ["Organization", "6 (invite, role, activate, deactivate)", "Admin only", "No"]])

h2("Service Layer (24 services)")
tbl(["Category", "Services", "Key Capabilities"],
    [["Statistical", "StatisticalAnalysisService, BayesianMethods, InterimAnalysis", "Cox PH, IPTW, KM, propensity scores, MI, MMRM, E-value, meta-analysis, O'Brien-Fleming"],
     ["CDISC", "AdamService, SDTMService", "ADSL/ADAE/ADTTE/ADLB generation, DM/AE/LB/VS/EX/DS, validation"],
     ["Documents", "DocumentGenerator, CSRGenerator, TFLGenerator, ADRGGenerator, DefineXMLGenerator, ECTDPackager", "SAR, CSR sections, demographics/AE/KM/forest/Love plots, ADRG, Define-XML 2.1, eCTD Module 5"],
     ["Evidence", "ExternalAPIs, SemanticScholar, AdvancedSearch", "PubMed/ClinTrials queries, semantic search, citation network"],
     ["Review", "CollaborativeReview, IntelligentWorkflow", "Assignments, comments, conflict resolution, workflow orchestration"],
     ["AI/ML", "EnhancedAI, LLMIntegration", "Evidence critique, entity extraction, quality assessment"],
     ["Infrastructure", "EmailService, TaskQueue, ProgramDashboard", "SMTP/console email, async tasks, portfolio analytics"]])

h2("Core Layer (10 modules)")
tbl(["Module", "File", "Responsibility"],
    [["Security", "security.py", "JWT (access + refresh + rotation), bcrypt, RBAC (4 roles), HTTPS redirect, CSP"],
     ["Database", "database.py", "Async SQLAlchemy (SQLite/PostgreSQL), connection pooling, health checks"],
     ["Config", "config.py", "170+ settings from .env, auto SECRET_KEY generation, production validation"],
     ["Cache", "cache.py", "Redis + in-memory LRU (2000 keys), TTL, pattern delete, @cached decorator"],
     ["Rate Limiter", "rate_limiter.py", "Redis + in-memory fallback (fail-closed), sliding window, per-endpoint limits"],
     ["Storage", "storage.py", "S3 + local filesystem, SHA-256 integrity, signed URLs, artifact persistence"],
     ["Observability", "observability.py", "Sentry init, RequestMetrics (p50/p95/p99), Server-Timing header"],
     ["Pagination", "pagination.py", "PaginationParams dependency, paginate_query(), standardized envelope"],
     ["Task Queue", "task_queue.py", "InProcessTaskQueue, TaskResult lifecycle, progress tracking"],
     ["Exceptions", "exceptions.py", "Custom exception hierarchy, standardized error responses, correlation IDs"]])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# DATA FLOW DIAGRAM
# ══════════════════════════════════════════════════════════════════
h1("Data Flow Diagram")
p("This diagram traces data from ingestion through processing to regulatory output.", italic=True, color=gray)

doc.add_paragraph()
code("""
  INGESTION                    PROCESSING                      OUTPUT
  =========                    ==========                      ======

  Protocol (PDF/DOCX)          Parsed Specification
       |                            |
       v                            v
  [Protocol Parser] ---------> processing_config.study_definition
       |                            |
       |                     +------+------+
       |                     |             |
       v                     v             v
  [Evidence Discovery]  [Cohort Builder]  [Covariate Config]
   PubMed API  ----+         |                  |
   ClinTrials  ----|         v                  v
   Sem. Scholar ---+--> EvidenceRecord    processing_config.cohort
                        table (n rows)    processing_config.covariates
                            |
                   +--------+--------+
                   |                 |
                   v                 v
            [Comparability     [ADaM Generator]
             Scoring]               |
                   |          +-----+-----+
                   v          |     |     |
            ComparabilityScore ADSL ADAE ADTTE
                   |          |     |     |
                   v          +-----+-----+
            [Bias Analysis]        |
                   |               v
                   v          [Statistical Engine]
            BiasAnalysis       Cox PH, IPTW, KM
                   |           Propensity Scores
                   |           MI, MMRM, Bayesian
                   |               |
                   v               v
            [Evidence Review]  [TFL Generator]
            ReviewDecision     Demographics Table
            (e-signed)         AE Table, KM Curves
                   |           Forest Plot, Love Plot
                   |               |
                   v               v
            [SAP Generator]    [CSR Generator]
                   |           Synopsis, Sec 11/12
                   |           Appendix 16
                   v               |
            [Define-XML Gen]       |
            [ADRG Generator]       |
                   |               |
                   +-------+-------+
                           |
                           v
                    [eCTD Packager]
                    m5/datasets/
                    m5/clinical-study-reports/
                    m5/literature-references/
                    Study Tagging File
                           |
                           v
                    FDA/EMA SUBMISSION
""")

h2("Data Flow Stages")
tbl(["Stage", "Input", "Process", "Output", "Storage"],
    [["1. Ingestion", "Protocol PDF/DOCX, research intent", "Text extraction + NLP parsing", "Structured specification", "Project.processing_config"],
     ["2. Discovery", "Search terms from specification", "PubMed/ClinTrials API queries", "Evidence records (n=50-100)", "evidence_records table"],
     ["3. Scoring", "Evidence + specification", "6-dimension comparability scoring", "Scores per evidence item", "comparability_scores table"],
     ["4. Bias Analysis", "Comparability scores", "5 bias type detection + E-value", "Bias severity + fragility", "bias_analyses table"],
     ["5. Dataset Generation", "Specification + cohort", "CDISC variable derivation", "ADaM (ADSL/ADAE/ADTTE)", "adam_datasets table + storage"],
     ["6. Statistical Analysis", "ADaM datasets", "Cox PH, IPTW, KM, MI, MMRM", "Effect estimates + CIs", "processing_config.results"],
     ["7. TFL Generation", "Statistical results + datasets", "Table/figure rendering", "Demographics, KM, forest plots", "regulatory_artifacts + storage"],
     ["8. Document Generation", "All above outputs", "Template assembly", "CSR, SAP, ADRG, Define-XML", "regulatory_artifacts + storage"],
     ["9. Submission Packaging", "All documents + datasets", "eCTD directory assembly", "eCTD Module 5 package", "storage (S3/local)"]])

h2("Data Persistence Model")
tbl(["Data Type", "Primary Store", "Secondary Store", "Retention"],
    [["User credentials", "PostgreSQL (users table)", "N/A", "Account lifetime"],
     ["Session tokens", "PostgreSQL (session_tokens)", "N/A", "7 days (auto-purge)"],
     ["Project config", "PostgreSQL (projects.processing_config JSON)", "Cache (30s TTL)", "Project lifetime"],
     ["Evidence records", "PostgreSQL (evidence_records)", "Cache (120s TTL)", "Project lifetime"],
     ["Statistical results", "PostgreSQL (processing_config)", "N/A", "Project lifetime"],
     ["Generated documents", "Storage (S3/local)", "PostgreSQL (regulatory_artifacts metadata)", "7 years (regulatory)"],
     ["Audit logs", "PostgreSQL (audit_logs)", "N/A", "7 years (21 CFR Part 11)"],
     ["Cache entries", "Redis / in-memory", "N/A", "TTL-based (30s-120s)"]])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# ERD
# ══════════════════════════════════════════════════════════════════
h1("Entity Relationship Diagram")
p("24 tables organized into 5 domain groups.", italic=True, color=gray)

h2("Domain 1: Identity & Multi-Tenancy")
code("""
  +----------------+       +----------------+
  | organizations  |       |     users      |
  |----------------|       |----------------|
  | id (PK)        |<------| organization_id|
  | name (UNIQUE)  |       | id (PK)        |
  | slug (UNIQUE)  |       | email (UNIQUE) |
  | is_active      |       | full_name      |
  | created_at     |       | role (ENUM)    |
  +----------------+       | hashed_password|
         |                 | is_active      |
         |                 +----------------+
         |                        |
         v                        v
  +----------------+       +------------------+
  |   projects     |       | session_tokens   |
  |----------------|       |------------------|
  | id (PK)        |       | id (PK)          |
  | organization_id|       | user_id (FK)     |
  | title          |       | token_hash       |
  | status (ENUM)  |       | token_type       |
  | created_by(FK) |       | code_hash        |
  | research_intent|       | expires_at       |
  | processing_cfg |       | is_revoked       |
  +----------------+       +------------------+
""")

h2("Domain 2: Evidence Pipeline")
code("""
  projects
      |
      | 1:N
      v
  +--------------------+      +------------------------+
  | evidence_records   |      | comparability_scores   |
  |--------------------|      |------------------------|
  | id (PK)            |<-----| evidence_record_id(FK) |
  | project_id (FK)    |      | id (PK)                |
  | source_type (ENUM) |      | population_similarity  |
  | source_id          |      | endpoint_alignment     |
  | title, abstract    |      | covariate_coverage     |
  | authors (JSON)     |      | temporal_alignment     |
  | journal, year      |      | evidence_quality       |
  | structured_data    |      | overall_score          |
  | extraction_conf    |      | regulatory_viability   |
  +--------------------+      +------------------------+
      |                              |
      | 1:N                          | 1:N
      v                              v
  +--------------------+      +--------------------+
  | evidence_embeddings|      | bias_analyses      |
  |--------------------|      |--------------------|
  | id (PK)            |      | id (PK)            |
  | evidence_id (FK)   |      | comp_score_id (FK) |
  | embedding_vector   |      | bias_type (ENUM)   |
  | embedding_model    |      | bias_severity      |
  +--------------------+      | fragility_score    |
                              | regulatory_risk    |
  +--------------------+      +--------------------+
  | citation_relations |
  |--------------------|
  | citing_evidence(FK)|
  | cited_evidence(FK) |
  | relationship_type  |
  | confidence_score   |
  +--------------------+
""")

h2("Domain 3: Review & Decisions")
code("""
  evidence_records          users
       |                      |
       | N:1                  | N:1
       v                      v
  +--------------------+  +--------------------+
  | review_decisions   |  | review_assignments |
  |--------------------|  |--------------------|
  | id (PK)            |  | id (PK)            |
  | project_id (FK)    |  | evidence_id (FK)   |
  | evidence_id (FK)   |  | reviewer_id (FK)   |
  | reviewer_id (FK)   |  | assigned_by (FK)   |
  | decision (ENUM)    |  | role, status       |
  | confidence_level   |  | due_date           |
  | rationale          |  +--------------------+
  | decided_at         |
  +--------------------+  +--------------------+
                          | review_comments    |
  +--------------------+  |--------------------|
  | evidence_critiques |  | id (PK)            |
  |--------------------|  | evidence_id (FK)   |
  | id (PK)            |  | author_id (FK)     |
  | project_id (FK)    |  | parent_id (FK self)|
  | overall_assessment |  | content            |
  | strengths (JSON)   |  | mentions (JSON)    |
  | weaknesses (JSON)  |  | resolved_by (FK)   |
  | fda_likelihood     |  +--------------------+
  +--------------------+
""")

h2("Domain 4: Regulatory Outputs")
code("""
  projects
      |
      | 1:N                    1:N
      +----------------+-------------------+
      v                v                   v
  +------------------+ +----------------+ +----------------+
  | regulatory_      | | adam_datasets   | | parsed_        |
  |   artifacts      | |----------------|  | specifications |
  |------------------| | id (PK)        | |----------------|
  | id (PK)          | | project_id(FK) | | id (PK)        |
  | project_id (FK)  | | dataset_name   | | project_id(FK) |
  | artifact_type    | | dataset_label  | | indication     |
  | title, format    | | variables(JSON)| | population_def |
  | file_path        | | records_count  | | primary_endpt  |
  | file_size        | | data_content   | | inclusion_crit |
  | checksum (SHA256)| | validation_stat| | sample_size    |
  | generated_by(FK) | +----------------+ +----------------+
  | regulatory_agency|
  +------------------+
""")

h2("Domain 5: System & Audit")
code("""
  +--------------------+  +--------------------+  +--------------------+
  | audit_logs         |  | workflow_steps     |  | user_presence      |
  |--------------------|  |--------------------|  |--------------------|
  | id (PK)            |  | id (PK)            |  | id (PK)            |
  | project_id (FK)    |  | workflow_id        |  | user_id (FK)       |
  | user_id (FK)       |  | step_type          |  | resource_type      |
  | action             |  | step_order         |  | resource_id        |
  | resource_type/id   |  | required_reviewers |  | activity           |
  | old_values (JSON)  |  | status             |  | cursor_position    |
  | new_values (JSON)  |  +--------------------+  +--------------------+
  | ip_address         |
  | regulatory_signif  |  +--------------------+  +--------------------+
  | retention_years(7) |  | saved_searches     |  | notification_sets  |
  +--------------------+  |--------------------|  |--------------------|
                          | id, user_id (FK)   |  | id, user_id (FK)   |
  +--------------------+  | name, query        |  | notification_type  |
  | federated_nodes    |  | search_type        |  | email_enabled      |
  |--------------------|  | filters (JSON)     |  | frequency          |
  | node_id (UNIQUE)   |  +--------------------+  +--------------------+
  | institution_name   |
  | endpoint_url       |  +--------------------+  +--------------------+
  | trust_score        |  | constraint_patterns|  | evidence_patterns  |
  +--------------------+  |--------------------|  |--------------------|
                          | pattern_name/type  |  | pattern_name       |
                          | pattern_logic(JSON)|  | indication_category|
                          | severity_weight    |  | regulatory_outcome |
                          | usage_count        |  | approval_likelihood|
                          +--------------------+  +--------------------+
""")

h2("Relationship Summary")
tbl(["Relationship", "Type", "Cardinality"],
    [["Organization -> Users", "One-to-Many", "1 org : N users"],
     ["Organization -> Projects", "One-to-Many", "1 org : N projects"],
     ["User -> Projects (created_by)", "One-to-Many", "1 user : N projects"],
     ["Project -> EvidenceRecords", "One-to-Many", "1 project : N evidence"],
     ["Project -> ParsedSpecifications", "One-to-Many", "1 project : N specs"],
     ["Project -> RegulatoryArtifacts", "One-to-Many", "1 project : N artifacts"],
     ["Project -> AdamDatasets", "One-to-Many", "1 project : N datasets"],
     ["Project -> AuditLogs", "One-to-Many", "1 project : N logs"],
     ["EvidenceRecord -> ComparabilityScores", "One-to-Many", "1 evidence : N scores"],
     ["ComparabilityScore -> BiasAnalyses", "One-to-Many", "1 score : N biases"],
     ["EvidenceRecord -> ReviewDecisions", "One-to-Many", "1 evidence : N decisions"],
     ["EvidenceRecord -> ReviewComments", "One-to-Many", "1 evidence : N comments"],
     ["ReviewComment -> ReviewComment (parent)", "Self-referential", "Thread nesting"],
     ["EvidenceRecord -> CitationRelationship", "Self-referential (2 FKs)", "Citation network"],
     ["User -> SessionTokens", "One-to-Many", "1 user : N tokens"],
     ["User -> ReviewAssignments", "One-to-Many", "1 user : N assignments"],
     ["User -> ReviewComments", "One-to-Many", "1 user : N comments"]])

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# ENUM REFERENCE
# ══════════════════════════════════════════════════════════════════
h1("Enum Reference")

tbl(["Enum", "Values", "Used By"],
    [["UserRole", "ADMIN, REVIEWER, ANALYST, VIEWER", "users.role"],
     ["ProjectStatus", "DRAFT, PROCESSING, REVIEW, COMPLETED, ARCHIVED", "projects.status"],
     ["EvidenceSourceType", "PUBMED, CLINICALTRIALS, UPLOADED_DOCUMENT, INSTITUTIONAL_DATA, FEDERATED_SOURCE", "evidence_records.source_type"],
     ["BiasType", "SELECTION_BIAS, CONFOUNDING, MEASUREMENT_BIAS, TEMPORAL_BIAS, PUBLICATION_BIAS", "bias_analyses.bias_type"],
     ["ReviewDecisionEnum", "ACCEPTED, REJECTED, DEFERRED, PENDING", "review_decisions.decision"]])

# Footer
doc.add_paragraph()
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("Afarensis Enterprise v2.1 \u2014 Synthetic Ascension \u2014 Confidential")
r.font.size = Pt(9); r.font.color.rgb = gray; r.font.name = "Calibri"

doc.save("Afarensis_Architecture_DataFlow_ERD.docx")
print("Generated: Afarensis_Architecture_DataFlow_ERD.docx")
