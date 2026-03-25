"""Generate the Afarensis End-to-End User Tutorial."""
import sys
sys.path.insert(0, "../backend")

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from datetime import datetime

doc = Document()

# ── Page Setup ───────────────────────────────────────────────────────
for section in doc.sections:
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# ── Colors ───────────────────────────────────────────────────────────
navy = RGBColor(0x1E, 0x3A, 0x5F)
blue = RGBColor(0x25, 0x63, 0xEB)
dark = RGBColor(0x1F, 0x2A, 0x36)
gray = RGBColor(0x6B, 0x72, 0x80)
white = RGBColor(0xFF, 0xFF, 0xFF)
green = RGBColor(0x05, 0x96, 0x69)

style = doc.styles["Normal"]
style.font.name = "Calibri"
style.font.size = Pt(10.5)
style.font.color.rgb = dark

def h1(text):
    h = doc.add_heading(text, level=1)
    for r in h.runs: r.font.color.rgb = navy; r.font.name = "Calibri"
    return h

def h2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.color.rgb = navy; r.font.name = "Calibri"
    return h

def h3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.color.rgb = blue; r.font.name = "Calibri"
    return h

def p(text, bold=False, italic=False, size=10.5, color=dark, after=6):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.size = Pt(size); run.font.color.rgb = color
    run.bold = bold; run.italic = italic; run.font.name = "Calibri"
    para.paragraph_format.space_after = Pt(after)
    return para

def tip(text):
    para = doc.add_paragraph()
    run = para.add_run("TIP: ")
    run.bold = True; run.font.color.rgb = green; run.font.name = "Calibri"; run.font.size = Pt(10)
    run2 = para.add_run(text)
    run2.font.size = Pt(10); run2.font.color.rgb = dark; run2.font.name = "Calibri"
    para.paragraph_format.space_after = Pt(4)
    return para

def note(text):
    para = doc.add_paragraph()
    run = para.add_run("NOTE: ")
    run.bold = True; run.font.color.rgb = blue; run.font.name = "Calibri"; run.font.size = Pt(10)
    run2 = para.add_run(text)
    run2.font.size = Pt(10); run2.font.color.rgb = dark; run2.font.name = "Calibri"
    return para

def bullet(text, bold_prefix=None):
    para = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = para.add_run(bold_prefix); r.bold = True; r.font.name = "Calibri"; r.font.size = Pt(10.5)
        r2 = para.add_run(text); r2.font.name = "Calibri"; r2.font.size = Pt(10.5)
    else:
        if para.runs:
            for r in para.runs: r.font.name = "Calibri"
        else:
            r = para.add_run(text); r.font.name = "Calibri"; r.font.size = Pt(10.5)
    return para

def step(number, text):
    para = doc.add_paragraph()
    r1 = para.add_run(f"Step {number}: ")
    r1.bold = True; r1.font.color.rgb = blue; r1.font.name = "Calibri"; r1.font.size = Pt(11)
    r2 = para.add_run(text)
    r2.font.name = "Calibri"; r2.font.size = Pt(10.5); r2.font.color.rgb = dark
    para.paragraph_format.space_after = Pt(4)
    return para

def code(text):
    para = doc.add_paragraph()
    run = para.add_run(text)
    run.font.name = "Consolas"; run.font.size = Pt(9); run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
    para.paragraph_format.space_after = Pt(6)
    return para

def add_table_row(table, cells, header=False):
    row = table.add_row()
    for i, text in enumerate(cells):
        cell = row.cells[i]
        cell.text = ""
        pp = cell.paragraphs[0]
        run = pp.add_run(str(text))
        run.font.size = Pt(9); run.font.name = "Calibri"
        if header:
            run.bold = True; run.font.color.rgb = white
            shading = cell._element.get_or_add_tcPr()
            bg = shading.makeelement(qn("w:shd"), {qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): "1E3A5F"})
            shading.append(bg)
        else:
            run.font.color.rgb = dark

def table(headers, rows):
    t = doc.add_table(rows=0, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    add_table_row(t, headers, header=True)
    for row in rows: add_table_row(t, row)
    doc.add_paragraph()
    return t


# ══════════════════════════════════════════════════════════════════════
# TITLE PAGE
# ══════════════════════════════════════════════════════════════════════
for _ in range(5): doc.add_paragraph()

pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("AFARENSIS"); r.font.size = Pt(36); r.font.color.rgb = navy; r.bold = True; r.font.name = "Calibri"

pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("End-to-End User Tutorial"); r.font.size = Pt(20); r.font.color.rgb = blue; r.font.name = "Calibri"

pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("From Login to FDA Submission Package"); r.font.size = Pt(13); r.font.color.rgb = gray; r.italic = True; r.font.name = "Calibri"

doc.add_paragraph(); doc.add_paragraph()

pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run(f"Version 2.1  |  {datetime.now().strftime('%B %Y')}  |  Synthetic Ascension")
r.font.size = Pt(10); r.font.color.rgb = gray; r.font.name = "Calibri"

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# TABLE OF CONTENTS (Manual)
# ══════════════════════════════════════════════════════════════════════
h1("Table of Contents")
toc_items = [
    "1. Getting Started: Login and Navigation",
    "2. Creating Your First Project",
    "3. Defining Your Study Protocol",
    "4. Discovering External Evidence",
    "5. Scoring Evidence Comparability",
    "6. Detecting and Quantifying Bias",
    "7. Running Statistical Analyses",
    "8. Handling Missing Data",
    "9. Generating CDISC Datasets (ADaM & SDTM)",
    "10. Creating Tables, Figures, and Listings",
    "11. Authoring the Statistical Analysis Plan",
    "12. Collaborative Evidence Review",
    "13. Generating Clinical Study Report Sections",
    "14. Preparing the Submission Package (eCTD)",
    "15. Monitoring and Administration",
    "Appendix A: Account Types and Permissions",
    "Appendix B: Keyboard Shortcuts",
    "Appendix C: Troubleshooting",
]
for item in toc_items:
    p(item, size=11)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 1: GETTING STARTED
# ══════════════════════════════════════════════════════════════════════
h1("1. Getting Started: Login and Navigation")

h2("1.1 Logging In")
p("Open your browser and navigate to your organization's Afarensis URL.")

step(1, "Enter your email address and password.")
step(2, "Click Sign In.")
step(3, "If this is your first login, you will be prompted to change your temporary password.")

note("If you enter the wrong password, you will see a red error banner. There is no lockout; you can retry immediately.")

h2("1.2 Forgot Password")
step(1, "Click Forgot password? on the login screen.")
step(2, "Enter your email address and click Send verification code.")
step(3, "Check your email for a 6-digit code. Enter it in the verification field.")
step(4, "Create a new password meeting the strength requirements (8+ characters, uppercase, lowercase, number, special character).")
step(5, "Click Reset password. You will be redirected to the login screen.")

tip("After a password reset, all your existing sessions are revoked for security. You will need to log in again on all devices.")

h2("1.3 The Dashboard")
p("After login, you land on the Dashboard. This is your command center.")

table(
    ["Section", "What It Shows"],
    [
        ["Projects", "Your active studies with status badges (Draft, Processing, Review, Completed)"],
        ["Quick Actions", "Create New Project, Generate SAR, Run Analysis"],
        ["System Health", "Database status, API availability, background task queue"],
    ],
)

h2("1.4 The Sidebar")
p("The left sidebar provides navigation through the 10-step regulatory workflow:")

table(
    ["Step", "Page", "Purpose"],
    [
        ["1", "Study Definition", "Define your protocol, endpoints, and population"],
        ["2", "Causal Framework", "Specify covariates and causal structure"],
        ["3", "Data Provenance", "Configure data sources and validate quality"],
        ["4", "Cohort Construction", "Set inclusion/exclusion criteria and run attrition"],
        ["5", "Comparability & Balance", "Assess covariate balance with Love plots"],
        ["6", "Bias & Sensitivity", "Run bias detection, E-value, and sensitivity analyses"],
        ["7", "Effect Estimation", "Execute primary analysis (Cox PH, IPTW, etc.)"],
        ["8", "Reproducibility", "Track computational environment and verify outputs"],
        ["9", "Audit Trail", "View immutable log of all system actions"],
        ["10", "Regulatory Output", "Generate and download submission documents"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 2: CREATING A PROJECT
# ══════════════════════════════════════════════════════════════════════
h1("2. Creating Your First Project")

step(1, "From the Dashboard, click Create New Project.")
step(2, "Enter a project title (e.g., \"XY-301: Cetuximab + FOLFIRI in mCRC\").")
step(3, "Add a description explaining the study purpose.")
step(4, "Enter the research intent. This drives evidence discovery. Be specific:")
code("  \"Evaluate the efficacy and safety of cetuximab combined with FOLFIRI\n   versus FOLFIRI alone in patients with KRAS wild-type metastatic\n   colorectal cancer, using overall survival as the primary endpoint.\"")
step(5, "Click Create. Your project is now in Draft status.")

tip("The research intent is the most important field. It determines what evidence the system discovers from PubMed and ClinicalTrials.gov. Write it as you would a research question.")

note("Projects are scoped to your organization. Other organizations cannot see your projects.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 3: STUDY DEFINITION
# ══════════════════════════════════════════════════════════════════════
h1("3. Defining Your Study Protocol")

p("Navigate to Step 1: Study Definition in the sidebar.")

h2("3.1 Protocol Specification")
step(1, "Fill in the study design fields:")
bullet("Indication (e.g., metastatic colorectal cancer)")
bullet("Population (e.g., adults with KRAS wild-type mCRC, ECOG 0-1)")
bullet("Intervention (e.g., Cetuximab 400mg/m2 loading + 250mg/m2 weekly + FOLFIRI)")
bullet("Comparator (e.g., FOLFIRI alone)")
bullet("Primary Endpoint (e.g., Overall Survival)")
bullet("Secondary Endpoints (add multiple)")

step(2, "Upload your protocol document (PDF or DOCX) using the Upload button. The system will extract structured fields automatically.")
step(3, "Review the parsed fields and correct any extraction errors.")
step(4, "Click Save to persist your study definition.")

h2("3.2 Estimand Framework (ICH E9(R1))")
p("Scroll down to the Estimand Framework section.")
step(1, "Define the primary estimand:")
bullet("Population: ", bold_prefix="Population: ")
bullet("Variable: ", bold_prefix="Variable: ")
bullet("Intercurrent Event Strategy: ", bold_prefix="ICE Strategy: ")
bullet("Summary Measure: ", bold_prefix="Summary Measure: ")
step(2, "Add secondary estimands as needed.")

h2("3.3 Locking the Protocol")
p("When your protocol is finalized:")
step(1, "Click Lock Protocol. This creates an immutable audit record.")
step(2, "After locking, the study definition becomes read-only. Changes require admin unlock with audit justification.")

tip("Lock the protocol before running analyses. This ensures your analysis results are traceable to a frozen specification.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 4: EVIDENCE DISCOVERY
# ══════════════════════════════════════════════════════════════════════
h1("4. Discovering External Evidence")

p("Navigate to Step 3: Data Provenance or use the Discover Evidence button on the project page.")

h2("4.1 Automatic Discovery")
step(1, "Click Discover Evidence. The system searches:")
bullet("PubMed (up to 50 results)")
bullet("ClinicalTrials.gov (up to 50 results)")
bullet("Semantic Scholar (optional)")

step(2, "This runs as a background task. You will see a progress indicator:")
bullet("10% - Searching PubMed...")
bullet("50% - Searching ClinicalTrials.gov...")
bullet("90% - Saving results...")
bullet("100% - Complete")

step(3, "When complete, review the discovered evidence on the Evidence tab.")

note("Evidence discovery is asynchronous. You can navigate away and come back; the task continues in the background. Check status via the Tasks indicator in the header.")

h2("4.2 Manual Upload")
step(1, "Click Upload Document on the project page.")
step(2, "Select a PDF, DOCX, or TXT file (max 100MB).")
step(3, "The system validates the file type using magic bytes, not just the file extension.")

h2("4.3 Reviewing Evidence")
p("Each evidence record shows:")
bullet("Title, authors, journal, year")
bullet("Source type (PubMed, ClinicalTrials, uploaded)")
bullet("Abstract preview")
bullet("Extraction confidence score (0-100%)")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 5: COMPARABILITY
# ══════════════════════════════════════════════════════════════════════
h1("5. Scoring Evidence Comparability")

p("Navigate to Step 5: Comparability & Balance.")

h2("5.1 Comparability Scoring")
p("For each evidence record, the system computes six dimensions:")

table(
    ["Dimension", "What It Measures", "Range"],
    [
        ["Population Similarity", "How closely the study population matches your protocol", "0-1"],
        ["Endpoint Alignment", "Whether the study measured your primary/secondary endpoints", "0-1"],
        ["Covariate Coverage", "How many of your specified covariates are reported", "0-1"],
        ["Temporal Alignment", "Whether follow-up periods and timeframes align", "0-1"],
        ["Evidence Quality", "Study design quality (RCT > observational > case series)", "0-1"],
        ["Provenance Score", "Data source reliability and provenance chain integrity", "0-1"],
    ],
)

step(1, "Review the overall comparability score (weighted average of all dimensions).")
step(2, "Click on any evidence row to see dimension-level detail.")

h2("5.2 Covariate Balance (Love Plot)")
step(1, "Click Compute Balance to calculate propensity scores and IPTW weights.")
step(2, "The Love Plot shows Standardized Mean Differences (SMD) for each covariate.")
step(3, "Covariates with |SMD| > 0.10 are flagged as imbalanced (red dots).")
step(4, "After weighting, the plot shows before/after balance improvement.")

tip("A well-balanced study should have all covariates below the 0.10 SMD threshold after weighting.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 6: BIAS DETECTION
# ══════════════════════════════════════════════════════════════════════
h1("6. Detecting and Quantifying Bias")

p("Navigate to Step 6: Bias & Sensitivity.")

h2("6.1 Automated Bias Detection")
step(1, "Click Run Bias Analysis. The system evaluates five bias types:")

table(
    ["Bias Type", "What It Detects"],
    [
        ["Selection Bias", "Non-random treatment assignment, selection by indication"],
        ["Confounding", "Unmeasured or inadequately controlled confounders"],
        ["Measurement Bias", "Inconsistent outcome ascertainment across groups"],
        ["Temporal Bias", "Immortal time bias, lead-time bias, time-window bias"],
        ["Publication Bias", "Over-representation of positive results in literature"],
    ],
)

step(2, "Review severity scores (0-1) for each bias type.")
step(3, "Review the E-value for unmeasured confounding:")
p("The E-value tells you how strong an unmeasured confounder would need to be to explain away your treatment effect. Higher E-values mean more robust results.", italic=True)

h2("6.2 Sensitivity Analyses")
p("The Sensitivity Catalog shows 24 pre-configured sensitivity analyses:")
bullet("Tipping-point analysis (how many outcomes would need to change to flip significance)")
bullet("Leave-one-out analysis (effect of removing each study)")
bullet("Subgroup analyses (by age, sex, region, etc.)")
bullet("Alternative estimand strategies")

h2("6.3 Missing Data Assessment")
step(1, "Click the Missing Data tab.")
step(2, "Review the missing data pattern summary.")
step(3, "Run Multiple Imputation (MICE) with Rubin's rules pooling.")
step(4, "Run Tipping-Point Analysis to find the delta where significance flips.")
step(5, "Run MMRM if you have repeated measures data.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 7: STATISTICAL ANALYSIS
# ══════════════════════════════════════════════════════════════════════
h1("7. Running Statistical Analyses")

p("Navigate to Step 7: Effect Estimation.")

h2("7.1 Primary Analysis")
p("The system supports these primary analysis methods:")

table(
    ["Method", "When to Use", "Output"],
    [
        ["Cox Proportional Hazards", "Time-to-event endpoints (OS, PFS)", "Hazard ratio, 95% CI, p-value"],
        ["IPTW (Inverse Probability Weighting)", "Observational data with confounding", "Weighted treatment effect"],
        ["Propensity Score Matching", "Creating balanced comparison groups", "Matched cohort analysis"],
        ["Kaplan-Meier", "Survival curve visualization", "Survival curves, median survival, log-rank p"],
    ],
)

step(1, "Select your primary analysis method.")
step(2, "Configure covariates for adjustment.")
step(3, "Click Run Analysis. Results appear in the forest plot.")
step(4, "Click Show Your Work on any result to see the full computational trace:")
bullet("Model Card: method specification, assumptions, limitations")
bullet("Formula: exact mathematical formulation")
bullet("Inputs: data sources and transformations")
bullet("Diagnostics: convergence, PH assumption test, goodness-of-fit")
bullet("Sensitivity: alternative specifications")
bullet("Lineage: full provenance chain from raw data to result")

h2("7.2 Multiplicity Adjustment")
p("If you have multiple endpoints:")
step(1, "Navigate to the Multiplicity panel.")
step(2, "Select your adjustment method (Bonferroni, Holm, Hochberg, or Benjamini-Hochberg).")
step(3, "View raw and adjusted p-values side by side.")
step(4, "Review the testing hierarchy visualization.")

h2("7.3 Bayesian Analysis (Optional)")
step(1, "Open the Bayesian panel.")
step(2, "Specify prior distributions (informative, weakly informative, or non-informative).")
step(3, "Run MCMC posterior sampling.")
step(4, "Review posterior distributions and credible intervals.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 8: MISSING DATA
# ══════════════════════════════════════════════════════════════════════
h1("8. Handling Missing Data")

p("Missing data handling is accessible from both Step 6 (Bias & Sensitivity) and Step 7 (Effect Estimation).")

h2("8.1 Multiple Imputation (MICE)")
step(1, "Navigate to the Missing Data tab.")
step(2, "Click Run Multiple Imputation.")
step(3, "Configure: number of imputations (default: 20), imputation model, variables to impute.")
step(4, "Review pooled results using Rubin's rules:")
bullet("Pooled estimate (average across imputations)")
bullet("Within-imputation variance")
bullet("Between-imputation variance")
bullet("Fraction of missing information")
bullet("Relative efficiency")

h2("8.2 Tipping-Point Analysis")
step(1, "Click Run Tipping-Point Analysis.")
step(2, "The system systematically shifts outcomes for missing subjects by increasing deltas.")
step(3, "Review the tipping delta: the value at which statistical significance reverses.")

h2("8.3 MMRM")
step(1, "Click Run MMRM (for repeated measures data).")
step(2, "Specify: outcome variable, time variable, group variable, subject ID, covariates.")
step(3, "Review fixed effects estimates, random effects, and fit statistics (AIC, BIC, log-likelihood).")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 9: CDISC DATASETS
# ══════════════════════════════════════════════════════════════════════
h1("9. Generating CDISC Datasets")

p("Navigate to Step 3: Data Provenance, then scroll to the CDISC section.")

h2("9.1 ADaM Datasets")
step(1, "Click Generate next to each ADaM dataset type:")

table(
    ["Dataset", "Full Name", "Content"],
    [
        ["ADSL", "Subject-Level Analysis Dataset", "One row per subject: demographics, treatment, population flags, dates"],
        ["ADAE", "Adverse Events Analysis Dataset", "One row per AE: MedDRA coding, severity, relationship, treatment-emergent flag"],
        ["ADTTE", "Time-to-Event Analysis Dataset", "One row per parameter per subject: time, censoring, event description"],
        ["ADLB", "Laboratory Analysis Dataset", "One row per lab parameter per visit: values, shift from baseline"],
    ],
)

step(2, "After generation, click Validate to check against ADaM Implementation Guide rules.")
step(3, "Review the validation report: variable names, labels, required variables, codelist compliance.")

h2("9.2 SDTM Datasets")
step(1, "Click Generate next to each SDTM domain:")
bullet("DM (Demographics), AE (Adverse Events), LB (Laboratory)")
bullet("VS (Vital Signs), EX (Exposure/Dosing), DS (Disposition)")
step(2, "Click Validate to check against SDTM IG rules.")
step(3, "Click Annotated CRF to view the CRF-to-SDTM mapping.")

note("CDISC dataset generation uses simulated data that matches your study definition (sample size, arms, endpoints). For production submissions, replace with real clinical data.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 10: TFL GENERATION
# ══════════════════════════════════════════════════════════════════════
h1("10. Creating Tables, Figures, and Listings")

p("Navigate to Step 10: Regulatory Output, then click the TFL section.")

h2("10.1 Available TFLs")

table(
    ["TFL Number", "Title", "Type"],
    [
        ["Table 14.1.1", "Demographics and Baseline Characteristics", "Table"],
        ["Table 14.3.1", "Adverse Events by System Organ Class", "Table"],
        ["Figure 14.1.1", "Covariate Balance Love Plot", "Figure"],
        ["Figure 14.2.1", "Kaplan-Meier Survival Curves", "Figure"],
        ["Figure 14.2.2", "Forest Plot of Treatment Effects", "Figure"],
    ],
)

step(1, "Click Generate next to any TFL to produce it.")
step(2, "Click Generate All to produce the complete TFL package.")
step(3, "Download individual TFLs or the full package as a ZIP.")

h2("10.2 Demographics Table")
p("The demographics table (Table 14.1.1) shows:")
bullet("Continuous variables: n, mean, SD, median, min, max")
bullet("Categorical variables: n, percentage")
bullet("Between-group comparison p-values")
bullet("Columns: Treatment, Control, Total")

h2("10.3 Kaplan-Meier Curves")
p("The KM curve (Figure 14.2.1) includes:")
bullet("Survival probability curves for each arm")
bullet("95% confidence bands")
bullet("At-risk table below the plot")
bullet("Median survival annotation")
bullet("Log-rank p-value")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 11: SAP
# ══════════════════════════════════════════════════════════════════════
h1("11. Authoring the Statistical Analysis Plan")

step(1, "Navigate to Step 10: Regulatory Output.")
step(2, "Click Generate SAP.")
step(3, "The system produces a 17-section ICH E9(R1)-compliant SAP document (DOCX) including:")

table(
    ["Section", "Content"],
    [
        ["1-2", "Title page, table of contents"],
        ["3-4", "Study objectives and design overview"],
        ["5", "Study populations (ITT, mITT, PP, Safety) with flag definitions"],
        ["6", "Estimand framework: population, variable, ICE strategies, summary measure"],
        ["7", "Endpoints with operational definitions"],
        ["8", "Sample size justification"],
        ["9", "Statistical methods: primary model, covariate adjustment, subgroup analyses"],
        ["10", "Missing data strategy: MMRM/MI primary, tipping-point sensitivity"],
        ["11", "Multiplicity adjustment: testing hierarchy and alpha allocation"],
        ["12", "Interim analysis plan (if applicable)"],
        ["13", "TFL shell specifications"],
        ["14-15", "Software and references"],
    ],
)

step(4, "Download the DOCX file and review.")

tip("Generate the SAP after defining your study protocol, covariates, and analysis methods. The SAP pulls from all configured sections.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 12: COLLABORATIVE REVIEW
# ══════════════════════════════════════════════════════════════════════
h1("12. Collaborative Evidence Review")

h2("12.1 Assigning Reviewers")
step(1, "From the project page, click Assign Reviewers.")
step(2, "Select evidence records to review.")
step(3, "Assign team members with a role (Reviewer, Senior Reviewer, Approver).")
step(4, "Set priority (Low, Normal, High, Urgent) and due dates.")

h2("12.2 Reviewing Evidence")
step(1, "Reviewers see their assignments on the Dashboard.")
step(2, "Click an assignment to open the evidence review page.")
step(3, "Read the evidence, comparability scores, and bias analysis.")
step(4, "Add comments (threaded, with @mentions for team members).")
step(5, "Submit a decision:")
bullet("Accepted: evidence is suitable for inclusion")
bullet("Rejected: evidence does not meet criteria")
bullet("Deferred: more information needed")
step(6, "Set a confidence level (0-100%) and provide written rationale.")
step(7, "Submit. Your decision is cryptographically signed (SHA-256 hash of decision + timestamp + user ID).")

h2("12.3 Conflict Resolution")
p("When reviewers disagree:")
step(1, "The system flags the conflict.")
step(2, "A senior reviewer or admin navigates to Review > Conflicts.")
step(3, "Reviews both decisions and rationales side by side.")
step(4, "Submits a resolution with justification.")

note("All review decisions, comments, and conflict resolutions are recorded in the immutable audit trail.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 13: CSR GENERATION
# ══════════════════════════════════════════════════════════════════════
h1("13. Generating Clinical Study Report Sections")

p("Navigate to Step 10: Regulatory Output.")

h2("13.1 Available CSR Sections")

table(
    ["Section", "Title", "Content"],
    [
        ["Synopsis", "Study Synopsis", "One-page summary of study design, methods, and key results"],
        ["Section 11", "Efficacy Evaluation", "Primary and secondary endpoint analyses, subgroup results"],
        ["Section 12", "Safety Evaluation", "AE tables, SAE analysis, deaths, discontinuations"],
        ["Appendix 16.1.9", "Statistical Methods", "Detailed description of all statistical methods used"],
    ],
)

step(1, "Click Generate next to any section, or Generate Full CSR for all sections at once.")
step(2, "Each section is saved as a DOCX file and stored as a regulatory artifact.")
step(3, "Download from the Artifacts list or from the project page.")

h2("13.2 ADRG (Analysis Data Reviewer's Guide)")
step(1, "Click Generate ADRG.")
step(2, "The system produces a PhUSE-compliant ADRG documenting:")
bullet("Computational environment (Python version, package list)")
bullet("Dataset structure and relationships")
bullet("Program listing and execution order")
bullet("Variable derivation rules")

h2("13.3 Define-XML")
step(1, "Click Generate Define-XML.")
step(2, "The system produces a CDISC Define-XML 2.1 document containing:")
bullet("All ADaM dataset descriptions")
bullet("Variable-level metadata (name, label, type, origin)")
bullet("Codelist definitions")
bullet("Value-level metadata")
step(3, "Click Validate to check against Define-XML schema rules.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 14: eCTD SUBMISSION
# ══════════════════════════════════════════════════════════════════════
h1("14. Preparing the Submission Package (eCTD)")

p("This is the final step: assembling everything into an FDA-ready eCTD Module 5 package.")

h2("14.1 Submission Readiness Check")
step(1, "Navigate to Step 10: Regulatory Output.")
step(2, "Click Submission Readiness. The system checks:")

table(
    ["Checkpoint", "What Is Verified"],
    [
        ["Protocol locked", "Study definition is frozen"],
        ["Evidence reviewed", "All evidence has reviewer decisions"],
        ["Bias assessed", "Bias analysis completed for all comparisons"],
        ["TFLs generated", "All required tables, figures, and listings exist"],
        ["CSR sections", "Synopsis, Section 11, Section 12, Appendix 16 generated"],
        ["ADaM datasets", "ADSL, ADAE, ADTTE validated"],
        ["Define-XML", "Generated and validated"],
        ["ADRG", "Generated"],
    ],
)

step(3, "Fix any gaps before proceeding.")

h2("14.2 Generating the eCTD Package")
step(1, "Click Generate eCTD Package.")
step(2, "The system assembles:")
bullet("m5/datasets/analysis/ (ADaM datasets + Define-XML)")
bullet("m5/datasets/tabulations/ (SDTM datasets)")
bullet("m5/clinical-study-reports/ (CSR sections)")
bullet("m5/literature-references/ (evidence bibliography)")
bullet("Study Tagging File (STF)")
bullet("Document checksums and manifest")
step(3, "Click Validate to check eCTD structure completeness.")
step(4, "Download the package.")

tip("Run the validation step before downloading. It catches missing files, broken references, and structural issues that the FDA gateway would reject.")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# CHAPTER 15: MONITORING
# ══════════════════════════════════════════════════════════════════════
h1("15. Monitoring and Administration")

h2("15.1 Audit Trail")
p("Navigate to Step 9: Audit Trail.")
p("Every action in the system is recorded with:")
bullet("Who (user ID, email)")
bullet("What (action type: created, updated, deleted, generated, reviewed)")
bullet("When (timestamp, UTC)")
bullet("Where (IP address, user agent)")
bullet("Old and new values (for modifications)")
p("Audit logs are retained for 7 years per regulatory requirements.")

h2("15.2 User Management (Admin Only)")
step(1, "Navigate to Organization > Users.")
step(2, "Invite new users: enter email, name, and role.")
step(3, "Change roles: click a user, select new role (Admin, Reviewer, Analyst, Viewer).")
step(4, "Deactivate users: click Deactivate to prevent login without deleting the account.")

h2("15.3 System Health")
p("Admin users can access:")
bullet("System Metrics: request throughput, latency percentiles (p50/p95/p99), error rates")
bullet("Cache Stats: hit ratios, key count, backend type")
bullet("Storage Stats: file count, total size, backend (local/S3)")
bullet("Background Tasks: running/completed/failed task list")

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# APPENDIX A: ROLES
# ══════════════════════════════════════════════════════════════════════
h1("Appendix A: Account Types and Permissions")

table(
    ["Role", "Can Do", "Cannot Do"],
    [
        ["Admin", "Everything: user management, system config, all analysis + review", "N/A"],
        ["Reviewer", "Review evidence, submit decisions, add comments, generate artifacts", "Manage users, change system settings"],
        ["Analyst", "Create projects, discover evidence, run analyses, generate TFLs/CSR", "Submit review decisions, manage users"],
        ["Viewer", "View projects, evidence, results, and artifacts (read-only)", "Create, modify, or delete anything"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# APPENDIX B: SHORTCUTS
# ══════════════════════════════════════════════════════════════════════
h1("Appendix B: Keyboard Shortcuts")

table(
    ["Shortcut", "Action"],
    [
        ["Ctrl + S", "Save current page"],
        ["Ctrl + N", "Create new project"],
        ["Ctrl + /", "Open search"],
        ["Esc", "Close modal or drawer"],
        ["Tab", "Navigate between form fields"],
    ],
)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════
# APPENDIX C: TROUBLESHOOTING
# ══════════════════════════════════════════════════════════════════════
h1("Appendix C: Troubleshooting")

table(
    ["Problem", "Cause", "Solution"],
    [
        ["Page shows 'Something went wrong'", "JavaScript error on one page", "Click Try Again or Go to Dashboard. Navigate back to the page."],
        ["Login returns no error message", "Network issue or server down", "Check browser console (F12). Verify the server URL is correct."],
        ["Evidence discovery stuck at 0%", "Background task pending", "Check Tasks list in the header. The task may be queued behind others."],
        ["Statistical analysis returns empty", "No data loaded for project", "Ensure evidence has been discovered and covariates defined."],
        ["File upload rejected", "Wrong file type or too large", "Allowed: PDF, DOCX, DOC, TXT, MD. Max size: 100MB."],
        ["Cannot see other users' projects", "Multi-tenancy isolation", "You can only see projects in your organization. Contact admin."],
        ["Rate limit error (429)", "Too many requests", "Wait 60 seconds and retry. Login is limited to 5 attempts/minute."],
    ],
)

doc.add_paragraph()
doc.add_paragraph()

# Footer
pp = doc.add_paragraph(); pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = pp.add_run("Afarensis Enterprise v2.1 \u2014 Synthetic Ascension \u2014 Confidential")
r.font.size = Pt(9); r.font.color.rgb = gray; r.font.name = "Calibri"

# ── Save ─────────────────────────────────────────────────────────────
output = "Afarensis_User_Tutorial.docx"
doc.save(output)
print(f"Generated: {output}")
