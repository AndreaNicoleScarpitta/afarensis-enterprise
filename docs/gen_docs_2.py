#!/usr/bin/env python3
"""Generate Afarensis Enterprise docs 4-6: Biostatisticians HowTo, Annotated Report, Architecture."""

import sys, os, json
from datetime import date

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
TODAY = date.today().strftime("%B %d, %Y")
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
BLUE = RGBColor(0x25, 0x63, 0xEB)

# ── Helpers (same as gen_docs_1) ────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def docx_title_page(doc, title, subtitle="", version="v2.1"):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AFARENSIS ENTERPRISE")
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(20)
    run2.font.color.rgb = NAVY
    run2.font.name = "Calibri"

    if subtitle:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(subtitle)
        run3.font.size = Pt(14)
        run3.font.color.rgb = BLUE
        run3.font.name = "Calibri"

    doc.add_paragraph("")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"Version {version}  |  {TODAY}")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run4.font.name = "Calibri"

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Synthetic Ascension Pty Ltd")
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run5.font.name = "Calibri"

    doc.add_page_break()

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def add_code_block(doc, text):
    """Add a monospaced code block paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    fmt = p.paragraph_format
    fmt.space_before = Pt(4)
    fmt.space_after = Pt(4)
    # light gray background via shading
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F5F5F5"/>')
    p._element.get_or_add_pPr().append(shading)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E3A5F")
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F0F4FA")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph("")
    return table

def add_callout(doc, title, body):
    """Add a callout box (bordered paragraph with bold title)."""
    p = doc.add_paragraph()
    # Add border via XML
    pPr = p._element.get_or_add_pPr()
    borders = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        '  <w:left w:val="single" w:sz="12" w:space="8" w:color="2563EB"/>'
        '  <w:top w:val="single" w:sz="4" w:space="2" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="2" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="2" w:color="CCCCCC"/>'
        '</w:pBdr>')
    pPr.append(borders)
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F0F7FF"/>')
    pPr.append(shading)

    run_t = p.add_run(f"[{title}] ")
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(10)
    run_t.font.color.rgb = BLUE

    run_b = p.add_run(body)
    run_b.font.name = "Calibri"
    run_b.font.size = Pt(10)
    return p


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENT 4: How-To for Biostatisticians
# ═════════════════════════════════════════════════════════════════════════════

def gen_howto():
    filepath = os.path.join(DOCS_DIR, "Afarensis_HowTo_Biostatisticians.docx")
    print(f"  Generating How-To for Biostatisticians...")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    docx_title_page(doc, "How-To Guide", "For Biostatisticians, HEOR, Regulatory & IT")

    # ── 1. Biostatisticians ──
    add_heading(doc, "1. For Biostatisticians", 1)

    add_heading(doc, "1.1 Cox Proportional Hazards Setup", 2)
    add_para(doc, "Navigate to Step 7 (Effect Estimation) and select Cox PH as the primary model. "
             "Configure the following parameters:")
    add_bullet(doc, "Time variable: Select the time-to-event column (e.g., OS_MONTHS, PFS_DAYS)")
    add_bullet(doc, "Event indicator: Select the censoring indicator (1 = event, 0 = censored)")
    add_bullet(doc, "Treatment variable: Binary treatment indicator from ADSL.TRTA")
    add_bullet(doc, "Covariates: Auto-populated from the causal framework adjustment set")
    add_para(doc, "The system automatically checks the proportional hazards assumption via Schoenfeld "
             "residuals and log-log plots. Violations are flagged with remediation suggestions: "
             "stratification, time-varying coefficients, or restricted mean survival time (RMST) as "
             "an alternative estimand.")

    add_heading(doc, "1.2 IPTW Configuration", 2)
    add_para(doc, "Inverse Probability of Treatment Weighting uses a logistic regression propensity "
             "score model by default. Configuration options:")
    add_table(doc, ["Parameter", "Default", "Options"], [
        ["PS Model", "Logistic Regression", "Logistic, GBM, Random Forest, Neural Net"],
        ["Weight Type", "ATE (unstabilized)", "ATE, ATT, ATC, stabilized ATE"],
        ["Trimming", "None", "Symmetric (1st/99th percentile), Crump optimal"],
        ["Overlap Assessment", "Enabled", "Propensity score density overlap plot"],
        ["Positivity Check", "Enabled", "Flags covariates with structural zeros"],
    ])

    add_heading(doc, "1.3 Propensity Score Methods", 2)
    add_para(doc, "Beyond IPTW, Afarensis supports:")
    add_bullet(doc, "PS Matching: 1:1 nearest-neighbor with caliper (0.2 * SD of logit PS)")
    add_bullet(doc, "PS Stratification: Quintile-based stratification with within-stratum analysis")
    add_bullet(doc, "Doubly Robust: Augmented IPTW (AIPTW) combining outcome model with PS weighting")
    add_bullet(doc, "Targeted Learning: TMLE with Super Learner ensemble for both PS and outcome models")

    add_heading(doc, "1.4 Missing Data Strategy", 2)
    add_para(doc, "Choose between two primary approaches based on the missing data mechanism:")
    add_table(doc, ["Method", "Assumption", "When to Use", "API Endpoint"], [
        ["Multiple Imputation (MI)", "MAR", "Baseline covariates with < 40% missing", "POST .../missing-data/impute"],
        ["MMRM", "MAR", "Longitudinal repeated measures with dropout", "POST .../missing-data/mmrm"],
        ["Tipping Point", "MNAR sensitivity", "After MI/MMRM to stress-test results", "POST .../missing-data/tipping"],
        ["Pattern-Mixture Models", "MNAR", "When MNAR is suspected clinically", "Via custom analysis code"],
    ])
    add_para(doc, "MI uses Rubin's rules to combine estimates across m=20 imputed datasets (configurable). "
             "The imputation model includes all analysis variables plus auxiliary variables predictive of "
             "missingness.")

    add_heading(doc, "1.5 Multiplicity Adjustment", 2)
    add_para(doc, "Select the appropriate method based on the regulatory context and endpoint structure:")
    add_table(doc, ["Method", "Controls", "Use Case"], [
        ["Bonferroni", "FWER", "Conservative; independent or correlated tests"],
        ["Holm (step-down)", "FWER", "Uniformly more powerful than Bonferroni"],
        ["Hochberg (step-up)", "FWER", "Requires independence or positive dependence"],
        ["Benjamini-Hochberg", "FDR", "Exploratory subgroup analyses"],
        ["Fixed-Sequence", "FWER", "Pre-specified testing hierarchy (primary, then key secondary)"],
        ["Graphical (Bretz)", "FWER", "Complex endpoint hierarchies with hypothesis recycling"],
    ])

    add_heading(doc, "1.6 Sensitivity Analysis Catalog", 2)
    add_para(doc, "The platform provides a pre-configured catalog of sensitivity analyses recommended "
             "by ICH E9(R1):")
    add_bullet(doc, "Per-Protocol Analysis: Excludes protocol violators")
    add_bullet(doc, "As-Treated Analysis: Classifies by actual treatment received")
    add_bullet(doc, "Tipping-Point Analysis: Determines how many events in the treatment arm would "
               "need to be reversed to nullify the result")
    add_bullet(doc, "E-Value Analysis: Quantifies robustness to unmeasured confounding")
    add_bullet(doc, "Rosenbaum Bounds: Tests sensitivity to hidden bias at various Gamma values")
    add_bullet(doc, "Negative Control Outcomes: Uses outcomes known to be unaffected by treatment as "
               "a falsification test")
    add_bullet(doc, "Quantitative Bias Analysis: Specifies bias parameters and adjusts estimates")

    add_heading(doc, "1.7 SAP Authoring Workflow", 2)
    add_para(doc, "The SAP generation endpoint (POST .../study/sap/generate) produces a Statistical "
             "Analysis Plan populated from the study definition and analysis configuration:")
    add_bullet(doc, "Section 1: Study Objectives and Endpoints (from study definition)")
    add_bullet(doc, "Section 2: Study Design (from study definition)")
    add_bullet(doc, "Section 3: Analysis Populations (ITT, mITT, PP, Safety)")
    add_bullet(doc, "Section 4: Statistical Methods (from effect estimation configuration)")
    add_bullet(doc, "Section 5: Missing Data Handling (from missing data strategy)")
    add_bullet(doc, "Section 6: Sensitivity Analyses (from sensitivity catalog selections)")
    add_bullet(doc, "Section 7: Subgroup Analyses (from pre-specified subgroups)")
    add_bullet(doc, "Section 8: Multiplicity Adjustment (from multiplicity configuration)")
    add_bullet(doc, "Section 9: TFL Shells (from TFL configuration)")

    # ── 2. HEOR ──
    add_heading(doc, "2. For HEOR Professionals", 1)

    add_heading(doc, "2.1 Comparative Effectiveness Setup", 2)
    add_para(doc, "Configure a comparative effectiveness research (CER) study by selecting 'Observational "
             "Cohort' as the study design. The system adapts the workflow to include target trial emulation "
             "parameters: eligibility criteria alignment, treatment strategy specification, outcome definition "
             "matching, and time-zero definition.")

    add_heading(doc, "2.2 External Control Arms", 2)
    add_para(doc, "To construct an external control arm from real-world data:")
    add_bullet(doc, "Define the target trial eligibility criteria in the cohort construction step")
    add_bullet(doc, "Import external data sources (CSV/SAS7BDAT) through the upload endpoint")
    add_bullet(doc, "The system matches external patients to trial inclusion/exclusion criteria")
    add_bullet(doc, "Propensity score adjustment accounts for confounding between trial and external arms")

    add_heading(doc, "2.3 Evidence Synthesis from Multiple Databases", 2)
    add_para(doc, "Afarensis searches five evidence databases simultaneously, then deduplicates results "
             "by DOI matching and title similarity (> 0.85 cosine similarity threshold). The merged evidence "
             "set preserves provenance metadata so you can trace each record back to its source API response.")

    add_heading(doc, "2.4 Publication Bias Assessment", 2)
    add_para(doc, "The bias analysis step includes publication bias assessment via:")
    add_bullet(doc, "Funnel plot asymmetry (visual inspection)")
    add_bullet(doc, "Egger's regression test for small-study effects")
    add_bullet(doc, "Trim-and-fill method for estimating missing studies")
    add_bullet(doc, "Selection model approaches (Vevea-Hedges weight function)")

    # ── 3. Regulatory Writers ──
    add_heading(doc, "3. For Regulatory Writers", 1)

    add_heading(doc, "3.1 CSR Section Mapping", 2)
    add_para(doc, "The CSR generator follows ICH E3 structure:")
    add_table(doc, ["ICH Section", "Content", "Afarensis Endpoint"], [
        ["Synopsis", "One-page study summary", "POST .../csr/synopsis"],
        ["Section 11", "Efficacy evaluation", "POST .../csr/section-11"],
        ["Section 12", "Safety evaluation", "POST .../csr/section-12"],
        ["Section 16", "Appendices (individual data)", "POST .../csr/appendix-16"],
        ["Full CSR", "All sections combined", "POST .../csr/full"],
    ])

    add_heading(doc, "3.2 ADRG Structure", 2)
    add_para(doc, "The Analysis Data Reviewer's Guide (generated via POST .../adrg/generate) includes:")
    add_bullet(doc, "Section 1: Introduction and study description")
    add_bullet(doc, "Section 2: Analysis datasets (ADSL, ADAE, ADTTE) with variable-level descriptions")
    add_bullet(doc, "Section 3: Data dependencies and derivation logic")
    add_bullet(doc, "Section 4: Computing environment (Python version, packages, OS)")
    add_bullet(doc, "Section 5: Analysis programs and execution order")
    add_bullet(doc, "Section 6: Issues summary and resolution log")

    add_heading(doc, "3.3 Define-XML Metadata", 2)
    add_para(doc, "Define-XML 2.1 is generated via POST .../define-xml/generate. Each ADaM variable entry "
             "includes: variable name, label (max 40 characters), data type, length, origin (CRF, Derived, "
             "Assigned, Protocol), computational method (for derived variables), codelist reference, and "
             "role (Identifier, Topic, Timing, Grouping, Record Qualifier, Result Qualifier, Synonym Qualifier).")

    add_heading(doc, "3.4 eCTD Module 5 Assembly", 2)
    add_para(doc, "The eCTD generator (POST .../ectd/generate) creates the Module 5 directory structure:")
    add_code_block(doc,
        "m5/\n"
        "  53-clin-stud-rep/\n"
        "    535-rep-analys-data-indiv-patient/\n"
        "      csr-synopsis.docx\n"
        "      csr-section-11-efficacy.docx\n"
        "      csr-section-12-safety.docx\n"
        "      csr-appendix-16.docx\n"
        "  datasets/\n"
        "    analysis/\n"
        "      adam/\n"
        "        adsl.xpt\n"
        "        adae.xpt\n"
        "        adtte.xpt\n"
        "        define.xml\n"
        "      adrg.docx\n"
        "    tabulation/\n"
        "      sdtm/\n"
        "        dm.xpt\n"
        "        ae.xpt\n"
        "        lb.xpt\n"
        "        vs.xpt\n"
        "        ex.xpt\n"
        "        ds.xpt\n"
        "  stf.xml (Study Tagging File)")

    add_heading(doc, "3.5 21 CFR Part 11 Compliance", 2)
    add_para(doc, "Afarensis supports 21 CFR Part 11 requirements through:")
    add_table(doc, ["Requirement", "Implementation"], [
        ["Electronic signatures", "SHA-256 signed decisions with user ID, timestamp, meaning"],
        ["Audit trail", "Immutable append-only log with 7-year retention"],
        ["Access controls", "RBAC with JWT authentication and session management"],
        ["System validation", "Documented IQ/OQ/PQ protocol with test evidence"],
        ["Record integrity", "Database-level constraints; hash verification on artifacts"],
        ["Backup and recovery", "Automated pg-backup to S3 with point-in-time recovery"],
    ])

    # ── 4. Data Managers ──
    add_heading(doc, "4. For Data Managers", 1)

    add_heading(doc, "4.1 CDISC ADaM Variable Mapping", 2)
    add_table(doc, ["Dataset", "Key Variables", "Derivation Notes"], [
        ["ADSL", "USUBJID, SITEID, AGE, AGEGR1, SEX, RACE, TRT01A, TRT01P, SAFFL, ITTFL, RANDDT",
         "Subject-level; one row per subject; flags derived from protocol criteria"],
        ["ADAE", "USUBJID, AETERM, AEDECOD (MedDRA PT), AEBODSYS, AESER, AEREL, AESTDTC, AEENDTC, TRTEMFL",
         "Treatment-emergent flag based on AESTDTC relative to first dose date"],
        ["ADTTE", "USUBJID, PARAMCD, PARAM, AVAL, CNSR, STARTDT, ADT, EVNTDESC, SRCDOM",
         "AVAL = ADT - STARTDT + 1; CNSR=0 for events, CNSR=1 for censored"],
    ])

    add_heading(doc, "4.2 SDTM Domain Generation", 2)
    add_para(doc, "Six SDTM domains are generated via POST .../sdtm/generate/{domain}:")
    add_table(doc, ["Domain", "Description", "Key Variables"], [
        ["DM", "Demographics", "STUDYID, USUBJID, SUBJID, SITEID, AGE, SEX, RACE, ARMCD, ARM"],
        ["AE", "Adverse Events", "AETERM, AEDECOD, AEBODSYS, AESEV, AESER, AEREL, AESTDTC, AEENDTC"],
        ["LB", "Laboratory", "LBTESTCD, LBTEST, LBORRES, LBORRESU, LBSTRESN, LBSTRESU, LBDTC"],
        ["VS", "Vital Signs", "VSTESTCD, VSTEST, VSORRES, VSORRESU, VSSTRESN, VSSTRESU, VSDTC"],
        ["EX", "Exposure", "EXTRT, EXDOSE, EXDOSU, EXROUTE, EXSTDTC, EXENDTC"],
        ["DS", "Disposition", "DSTERM, DSDECOD, DSCAT, DSSTDTC, EPOCH"],
    ])

    add_heading(doc, "4.3 Data Validation Rules", 2)
    add_para(doc, "The validate endpoints run CDISC conformance checks including:")
    add_bullet(doc, "Variable naming: Must match CDISC controlled terminology")
    add_bullet(doc, "Required variables: Checks presence of all required columns per domain")
    add_bullet(doc, "Value-level metadata: Validates codelists against CDISC dictionaries")
    add_bullet(doc, "Cross-domain consistency: USUBJID must exist in DM for all observation domains")
    add_bullet(doc, "Date formats: ISO 8601 compliance for all date/datetime variables")

    add_heading(doc, "4.4 Dataset Lineage", 2)
    add_para(doc, "The ADRG documents complete data lineage: raw data source, SDTM mapping, ADaM derivation "
             "logic, and TFL consumption. Each transformation step records the program name, version, "
             "execution timestamp, and input/output checksums.")

    # ── 5. IT/DevOps ──
    add_heading(doc, "5. For IT/DevOps", 1)

    add_heading(doc, "5.1 Deployment Architecture", 2)
    add_para(doc, "Afarensis deploys as a containerized stack:")
    add_table(doc, ["Component", "Technology", "Port", "Notes"], [
        ["Frontend", "React 18 SPA", "3000", "Served via Nginx in production"],
        ["API Server", "FastAPI (Python 3.10+)", "8000", "Uvicorn with auto-reload in dev"],
        ["Database", "PostgreSQL 15+", "5432", "With pgvector extension for embeddings"],
        ["Connection Pool", "PgBouncer", "6432", "Transaction-mode pooling, max 100 connections"],
        ["Cache", "Redis 7+", "6379", "Session store, task queue, caching"],
        ["Object Storage", "S3 / MinIO", "9000", "Artifacts, uploads, generated documents"],
        ["Error Tracking", "Sentry", "N/A", "DSN configured via SENTRY_DSN env var"],
        ["Backup", "pg-backup-s3", "N/A", "Hourly incremental, daily full, 30-day retention"],
    ])

    add_heading(doc, "5.2 Docker Deployment", 2)
    add_code_block(doc,
        "# docker-compose.yml (simplified)\n"
        "services:\n"
        "  api:\n"
        "    build: ./backend\n"
        "    ports: [\"8000:8000\"]\n"
        "    env_file: .env\n"
        "    depends_on: [db, redis]\n"
        "  db:\n"
        "    image: pgvector/pgvector:pg15\n"
        "    volumes: [pgdata:/var/lib/postgresql/data]\n"
        "  redis:\n"
        "    image: redis:7-alpine\n"
        "  pgbouncer:\n"
        "    image: edoburu/pgbouncer\n"
        "    depends_on: [db]")

    add_heading(doc, "5.3 API Integration", 2)
    add_para(doc, "All API endpoints are documented in _api_docs.json (146 endpoints across 18 categories). "
             "Authentication uses JWT bearer tokens obtained via POST /api/v1/auth/login. Tokens expire after "
             "15 minutes; refresh tokens last 7 days with automatic rotation.")

    add_heading(doc, "5.4 Webhook Setup", 2)
    add_para(doc, "Configure webhooks for event notifications:")
    add_bullet(doc, "project.created / project.completed: Notify external systems of project lifecycle events")
    add_bullet(doc, "evidence.discovered: Trigger downstream pipelines when new evidence is found")
    add_bullet(doc, "artifact.generated: Notify document management systems of new regulatory artifacts")
    add_bullet(doc, "review.decision: Integration with e-signature platforms for approval workflows")

    add_heading(doc, "5.5 Monitoring", 2)
    add_para(doc, "Key monitoring endpoints:")
    add_table(doc, ["Endpoint", "Purpose", "Frequency"], [
        ["GET /api/v1/health", "Basic liveness probe", "Every 10s (Kubernetes)"],
        ["GET /api/v1/system/health/detailed", "Component-level status", "Every 30s"],
        ["GET /api/v1/system/metrics", "Request latency, error rates", "Every 60s (Prometheus scrape)"],
        ["GET /api/v1/system/cache-stats", "Redis hit/miss ratio", "Every 60s"],
        ["GET /api/v1/system/storage-stats", "S3 usage and file counts", "Every 5m"],
    ])

    doc.save(filepath)
    sz = os.path.getsize(filepath)
    print(f"  -> {filepath} ({sz:,} bytes)")


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENT 5: Annotated Report
# ═════════════════════════════════════════════════════════════════════════════

def gen_annotated_report():
    filepath = os.path.join(DOCS_DIR, "Afarensis_Annotated_Report.docx")
    print(f"  Generating Annotated Report...")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    docx_title_page(doc, "Annotated Example Report", "Understanding Generated SAR/CSR Output")

    add_heading(doc, "About This Document", 1)
    add_para(doc, "This document walks through an example generated Study Analysis Report (SAR) / Clinical "
             "Study Report (CSR), with annotated callout boxes explaining why each section exists, what "
             "regulatory reviewers look for, and how Afarensis populates the content.")

    # ── Title Page ──
    add_heading(doc, "1. Title Page", 1)
    add_para(doc, "The title page of a CSR contains the study identification metadata required by ICH E3.")
    add_table(doc, ["Field", "Example Value", "Purpose"], [
        ["Study ID", "XY-301-2024", "Unique identifier; matches registry entries"],
        ["Protocol Version", "3.0 (Amendment 2)", "Tracks which protocol version was executed"],
        ["Report Date", "March 2026", "Date of report finalization"],
        ["Sponsor", "Synthetic Ascension Pty Ltd", "Legal entity responsible for the study"],
        ["Regulatory Submission", "NDA 123456", "Links to the broader submission dossier"],
    ])
    add_callout(doc, "Why This Matters",
                "FDA reviewers use the Study ID to cross-reference the CSR with the Study Tagging File "
                "in the eCTD. The protocol version confirms which study conduct rules were followed. "
                "Discrepancies between the title page and eCTD metadata trigger review queries.")

    # ── Synopsis ──
    add_heading(doc, "2. Synopsis", 1)
    add_para(doc, "The synopsis is a 1-2 page self-contained summary of the entire study. It must include: "
             "study objectives, design, key eligibility criteria, treatment arms, primary and key secondary "
             "endpoints, statistical methods, and top-line results.")
    add_callout(doc, "Regulatory Significance",
                "FDA and EMA reviewers read the synopsis first to decide whether a full review is warranted. "
                "The synopsis must stand alone without requiring reference to other sections. Every number "
                "in the synopsis must match the corresponding number in the body of the report.")

    # ── Section 11 ──
    add_heading(doc, "3. Section 11: Efficacy Evaluation", 1)
    add_para(doc, "Section 11 presents the primary efficacy analysis and all pre-specified secondary and "
             "sensitivity analyses.")

    add_heading(doc, "3.1 Primary Analysis (Cox PH)", 2)
    add_para(doc, "The primary result is presented as a Hazard Ratio with 95% confidence interval and "
             "two-sided p-value. Example: HR = 0.72 (95% CI: 0.58-0.89), p = 0.003.")
    add_callout(doc, "Interpretation",
                "A hazard ratio of 0.72 means the treatment arm has a 28% lower instantaneous risk of "
                "the event compared to the comparator at any given time point, assuming proportional hazards "
                "hold. The confidence interval excludes 1.0, and the p-value is below the pre-specified "
                "alpha of 0.05, indicating statistical significance.")

    add_heading(doc, "3.2 Forest Plot", 2)
    add_para(doc, "The forest plot shows subgroup analyses and sensitivity analyses alongside the primary result.")
    add_callout(doc, "How to Read a Forest Plot",
                "Diamond = overall summary estimate; its width spans the confidence interval. "
                "Squares = individual subgroup or analysis estimates; larger squares indicate greater weight. "
                "Horizontal lines = confidence intervals for each estimate. "
                "Vertical dashed line at HR=1.0 = null effect (no difference between arms). "
                "Estimates to the left of the line favor treatment; to the right favor comparator.")

    add_heading(doc, "3.3 Multiplicity Adjustment", 2)
    add_para(doc, "When multiple endpoints are tested, p-values are adjusted to control the family-wise "
             "error rate (FWER). The SAP pre-specifies the testing hierarchy and adjustment method.")
    add_callout(doc, "Why It Matters",
                "Without multiplicity adjustment, testing 20 endpoints at alpha=0.05 yields ~64% chance "
                "of at least one false positive. Regulatory agencies require adjustment to maintain the "
                "overall Type I error rate at the specified alpha level.")

    # ── Section 12 ──
    add_heading(doc, "4. Section 12: Safety Evaluation", 1)

    add_heading(doc, "4.1 Adverse Events Table Structure", 2)
    add_para(doc, "AEs are tabulated by System Organ Class (SOC) using MedDRA coding. The table shows "
             "the number and percentage of subjects experiencing each AE in each treatment arm.")
    add_table(doc, ["Column", "Content", "Purpose"], [
        ["SOC / Preferred Term", "MedDRA hierarchy", "Standardized classification for cross-study comparison"],
        ["Treatment Arm n (%)", "Event count and percentage", "Frequency in the active treatment group"],
        ["Comparator Arm n (%)", "Event count and percentage", "Frequency in the control group"],
        ["Risk Difference (95% CI)", "Absolute difference", "Clinical magnitude of the safety signal"],
    ])
    add_callout(doc, "Why Serious AEs Are Separated",
                "Serious adverse events (SAEs) are reported separately because they have distinct regulatory "
                "implications: expedited reporting to IRBs and regulatory agencies, potential protocol amendments, "
                "and potential labeling changes. SAEs include death, hospitalization, life-threatening events, "
                "disability, and congenital anomalies.")

    add_heading(doc, "4.2 Treatment-Emergent Definition", 2)
    add_callout(doc, "What Treatment-Emergent Means",
                "A treatment-emergent adverse event (TEAE) is one that first occurs or worsens in severity "
                "after the first dose of study treatment. Events present at baseline that do not worsen are "
                "not treatment-emergent. The TRTEMFL flag in ADAE implements this logic by comparing AESTDTC "
                "to the first dose date from ADSL.")

    # ── Demographics Table ──
    add_heading(doc, "5. Demographics Table", 1)
    add_para(doc, "Table 14.1.1 displays baseline characteristics for both treatment arms and the total "
             "population. It includes continuous variables (mean, SD, median, range) and categorical variables "
             "(n, percentage).")
    add_callout(doc, "Why Both Arms + Total",
                "Showing both arms allows assessment of balance between groups. The total column provides "
                "the overall study population characteristics. FDA reviewers check whether the study population "
                "matches the target indication population and whether randomization achieved balance.")
    add_callout(doc, "SMD > 0.10",
                "Standardized Mean Differences greater than 0.10 indicate meaningful imbalance between arms. "
                "While p-values test statistical significance of differences, SMD measures practical significance. "
                "Austin (2011) established 0.10 as the threshold for declaring balance in propensity score analyses.")

    # ── Kaplan-Meier ──
    add_heading(doc, "6. Kaplan-Meier Survival Curves", 1)
    add_para(doc, "Figure 14.2.1 shows the estimated survival probability over time for each treatment arm.")
    add_callout(doc, "At-Risk Table",
                "The at-risk table below the curve shows the number of subjects still under observation "
                "(not yet censored or experienced the event) at each time point. Reviewers use this to "
                "assess whether the curves are reliable at later time points (fewer subjects = wider CIs).")
    add_callout(doc, "Median Survival Annotation",
                "The median survival time is the time at which the Kaplan-Meier estimate crosses 50%. "
                "If the curve never crosses 50%, median survival is not estimable (NE). The annotation "
                "includes the 95% CI for the median, computed using the Brookmeyer-Crowley method.")
    add_callout(doc, "Log-Rank Test",
                "The log-rank test compares the survival distributions between groups. It tests the null "
                "hypothesis that the survival curves are identical. The p-value is reported on the figure. "
                "Note: log-rank is a non-parametric test that does not assume proportional hazards.")

    # ── Forest Plot Detail ──
    add_heading(doc, "7. Forest Plot (Detailed)", 1)
    add_table(doc, ["Element", "Visual", "Meaning"], [
        ["Diamond", "Wider shape at summary row", "Overall pooled estimate; width = confidence interval"],
        ["Squares", "Filled boxes per analysis", "Point estimate; size proportional to weight/sample size"],
        ["Horizontal lines", "Lines extending from squares", "95% confidence interval for that estimate"],
        ["Vertical line at 1.0", "Dashed reference line", "Null effect; no difference between arms"],
        ["Heterogeneity I-squared", "Footnote or header", "Percentage of variability due to true differences vs. chance"],
        ["Favors labels", "Below the x-axis", "Direction of effect: left = favors treatment, right = favors comparator"],
    ])

    # ── ADRG ──
    add_heading(doc, "8. ADRG (Analysis Data Reviewer's Guide)", 1)
    add_callout(doc, "Why Document the Computing Environment",
                "FDA reviewers may attempt to reproduce analyses. The ADRG must specify the exact software "
                "versions, operating system, and hardware configuration used. If the reviewer cannot reproduce "
                "the result, it may trigger a Complete Response Letter or Refuse-to-File action.")
    add_callout(doc, "What Reproducibility Means for FDA",
                "Reproducibility in the regulatory context means that an independent analyst, given the same "
                "data and programs, produces the same numerical results. This requires: deterministic algorithms, "
                "pinned package versions, documented random seeds, and archived execution environments.")

    # ── Define-XML ──
    add_heading(doc, "9. Define-XML", 1)
    add_callout(doc, "Why Every Variable Needs a Label, Origin, and Codelist",
                "Define-XML serves as the machine-readable data dictionary for the submission. Without it, "
                "FDA's automated review tools cannot parse the datasets. Every variable must have: "
                "a label (human-readable description, max 40 chars), an origin (CRF, Derived, Assigned, Protocol) "
                "documenting where the value comes from, and a codelist reference for categorical variables to "
                "ensure controlled terminology compliance.")

    # ── eCTD ──
    add_heading(doc, "10. eCTD Structure", 1)
    add_callout(doc, "Why m5/ Has Specific Subdirectories",
                "The eCTD (electronic Common Technical Document) is an internationally agreed format (ICH M8). "
                "Module 5 contains clinical study reports and is subdivided into: 5.3.5 (reports of analyses of "
                "data from more than one study), datasets/ (SDTM and ADaM), and supporting documents. The rigid "
                "directory structure allows regulatory agencies automated ingestion into their review systems.")
    add_callout(doc, "Study Tagging File (STF)",
                "The stf.xml file links the eCTD leaf documents to their regulatory context: which study, which "
                "indication, which submission sequence. It enables FDA's DERS (Document and Electronic Review System) "
                "to index and route the submission to the correct review division.")

    doc.save(filepath)
    sz = os.path.getsize(filepath)
    print(f"  -> {filepath} ({sz:,} bytes)")


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENT 6: Architecture
# ═════════════════════════════════════════════════════════════════════════════

def gen_architecture():
    filepath = os.path.join(DOCS_DIR, "Afarensis_Architecture_v2.1.docx")
    print(f"  Generating Architecture Document...")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    docx_title_page(doc, "Technical Architecture", "C4 Model, Data Flow, ERD & Dependencies")

    # ═══ C1 System Context ═══
    add_heading(doc, "1. C1: System Context Diagram", 1)
    add_para(doc, "The system context shows Afarensis Enterprise and its relationships with users and external systems.")
    add_code_block(doc,
        "+-------------------+     +-------------------+     +-------------------+\n"
        "|  Biostatistician  |     |     Reviewer       |     |      Admin        |\n"
        "|  (Primary User)   |     |  (Evidence Review) |     |  (System Config)  |\n"
        "+--------+----------+     +--------+----------+     +--------+----------+\n"
        "         |                         |                          |\n"
        "         +------------+------------+-----------+--------------+\n"
        "                      |                        |\n"
        "                      v                        v\n"
        "         +----------------------------------------------+\n"
        "         |       AFARENSIS ENTERPRISE v2.1               |\n"
        "         |  Evidence Synthesis & Regulatory Platform      |\n"
        "         |  146 API Endpoints | 10-Step Workflow          |\n"
        "         +------+-------+-------+-------+-------+--------+\n"
        "                |       |       |       |       |\n"
        "         +------+  +----+  +----+  +----+  +---+------+\n"
        "         v         v       v       v       v          v\n"
        "    +---------+ +------+ +------+ +-----+ +------+ +--------+\n"
        "    | PubMed  | |CT.gov| |OpenAl| |  SS | |BioGPT| |Claude/ |\n"
        "    | (NCBI)  | |      | |  ex  | |     | |(local| |GPT-4   |\n"
        "    +---------+ +------+ +------+ +-----+ +------+ +--------+\n"
        "\n"
        "    +----------+\n"
        "    | SendGrid |  (Email notifications)\n"
        "    +----------+\n"
        "\n"
        "    +----------+\n"
        "    | Regulator|  (Receives eCTD packages)\n"
        "    +----------+")

    add_table(doc, ["Actor / System", "Type", "Interaction"], [
        ["Biostatistician", "User", "Creates projects, configures studies, runs analyses, generates TFLs"],
        ["Reviewer", "User", "Reviews evidence, submits e-signed decisions, resolves conflicts"],
        ["Admin", "User", "Manages users, orgs, system config, monitors health"],
        ["Regulator", "External User", "Receives and reviews eCTD submission packages"],
        ["PubMed (NCBI)", "External System", "E-utilities API for biomedical literature search"],
        ["ClinicalTrials.gov", "External System", "v2 API for registered clinical trial search"],
        ["OpenAlex", "External System", "REST API for open-access scholarly works"],
        ["Semantic Scholar", "External System", "Academic Graph API for citation-aware paper search"],
        ["BioGPT", "Internal Model", "Local biomedical language model for text generation"],
        ["Claude / GPT-4", "External AI", "LLM for regulatory critique and evidence analysis"],
        ["SendGrid", "External Service", "Transactional email delivery for notifications"],
    ])

    # ═══ C2 Container ═══
    add_heading(doc, "2. C2: Container Diagram", 1)
    add_para(doc, "The container diagram shows the major deployable units and their communication.")
    add_code_block(doc,
        "                    +-------------------+\n"
        "                    |   React SPA       |\n"
        "                    |   Port 3000       |\n"
        "                    |   (Nginx in prod) |\n"
        "                    +--------+----------+\n"
        "                             |\n"
        "                             | HTTPS / JSON\n"
        "                             v\n"
        "                    +-------------------+\n"
        "                    |   FastAPI Server   |\n"
        "                    |   Port 8000        |\n"
        "                    |   (Uvicorn)        |\n"
        "                    +--+-----+-----+----+\n"
        "                       |     |     |\n"
        "              +--------+  +--+--+  +--------+\n"
        "              v           v     v           v\n"
        "    +----------------+ +-----+ +-------+ +--------+\n"
        "    | PostgreSQL 15  | |Redis| |  S3   | | Sentry |\n"
        "    | Port 5432      | |6379 | | 9000  | | (SaaS) |\n"
        "    | + pgvector     | |     | |       | |        |\n"
        "    +-------+--------+ +-----+ +-------+ +--------+\n"
        "            |\n"
        "    +-------+--------+\n"
        "    | PgBouncer      |\n"
        "    | Port 6432      |\n"
        "    +-------+--------+\n"
        "            |\n"
        "    +-------+--------+\n"
        "    | pg-backup-s3   |\n"
        "    | (cron job)     |\n"
        "    +----------------+")

    add_table(doc, ["Container", "Technology", "Port", "Responsibility"], [
        ["React SPA", "React 18, TypeScript, TailwindCSS", "3000", "User interface, 10-step workflow, data visualization"],
        ["FastAPI Server", "Python 3.10+, Uvicorn, Pydantic", "8000", "REST API (146 endpoints), business logic, auth"],
        ["PostgreSQL", "PostgreSQL 15 + pgvector", "5432", "Primary data store, embeddings, audit logs"],
        ["PgBouncer", "PgBouncer 1.21", "6432", "Connection pooling (transaction mode, max 100)"],
        ["Redis", "Redis 7 Alpine", "6379", "Session store, task queue, response caching"],
        ["S3 / MinIO", "S3-compatible storage", "9000", "Artifact storage, uploads, generated documents"],
        ["Sentry", "Sentry SaaS", "N/A", "Error tracking, performance monitoring"],
        ["pg-backup-s3", "Custom cron container", "N/A", "Automated database backups to S3"],
    ])

    # ═══ C3 Component ═══
    add_heading(doc, "3. C3: Component Diagram", 1)
    add_para(doc, "The API server is organized into three layers: API, Service, and Core.")

    add_heading(doc, "3.1 API Layer (146 Endpoints in 18 Groups)", 2)
    add_code_block(doc,
        "+------------------------------------------------------------------+\n"
        "|                        API LAYER                                 |\n"
        "|  +--------+ +--------+ +--------+ +--------+ +----------+       |\n"
        "|  | Auth   | |Projects| |Evidence| |Review  | |Study Cfg |       |\n"
        "|  | (9)    | | (4)    | | (8)    | | (10)   | | (16)     |       |\n"
        "|  +--------+ +--------+ +--------+ +--------+ +----------+       |\n"
        "|  +--------+ +--------+ +--------+ +--------+ +----------+       |\n"
        "|  | CDISC  | | TFL    | |Regulat.| |Search  | |Statistics|       |\n"
        "|  | (4)    | | (7)    | | (14)   | | (8)    | | (6)      |       |\n"
        "|  +--------+ +--------+ +--------+ +--------+ +----------+       |\n"
        "|  +--------+ +--------+ +--------+ +--------+ +----------+       |\n"
        "|  |Bayesian| |Interim | | SDTM   | |Program | | System   |       |\n"
        "|  | (3)    | | (3)    | | (4)    | | (4)    | | (8)      |       |\n"
        "|  +--------+ +--------+ +--------+ +--------+ +----------+       |\n"
        "|  +--------+ +--------+ +--------+                               |\n"
        "|  | BioGPT | | Tasks  | |Org Mgt |                               |\n"
        "|  | (4)    | | (4)    | | (6)    |                               |\n"
        "|  +--------+ +--------+ +--------+                               |\n"
        "+------------------------------------------------------------------+")

    add_table(doc, ["Group", "# Endpoints", "Key Responsibilities"], [
        ["Auth", "9", "Login, logout, JWT refresh, password reset, session management"],
        ["Projects", "4", "CRUD operations, protocol upload, project listing"],
        ["Evidence", "8", "Discovery, comparability, bias analysis, critique, decisions"],
        ["Review", "10", "Workflows, assignments, comments, conflicts, presence"],
        ["Study Config", "16", "Definition, covariates, data sources, cohort, balance, bias, reproducibility, audit"],
        ["CDISC (ADaM)", "4", "Generate, list, validate, metadata for ADaM datasets"],
        ["TFL", "7", "Demographics, AE table, KM curve, forest plot, Love plot, shells, generate-all"],
        ["Regulatory", "14", "SAR pipeline, CSR sections, SAP, ADRG, Define-XML, eCTD"],
        ["Search", "8", "Semantic, hybrid, recommendations, saved, citation network, Semantic Scholar"],
        ["Statistics", "6", "Full analysis, summary, MI, MMRM, tipping point, missing data summary"],
        ["Bayesian", "3", "Full analysis, prior elicitation, adaptive decisions"],
        ["Interim", "3", "Boundaries, evaluation, DSMB report"],
        ["SDTM", "4", "Generate domain, generate all, validate, aCRF"],
        ["Program", "4", "Overview, portfolio, readiness, milestones"],
        ["System", "8", "Health, detailed health, metrics, cache stats, storage, analytics, audit, data classify"],
        ["BioGPT", "4", "Status, generate, explain mechanism, summarize"],
        ["Tasks", "4", "Status, result, list, cancel"],
        ["Org Management", "6", "Info, users list, invite, role update, activate, deactivate"],
    ])

    add_heading(doc, "3.2 Service Layer (24 Services in 7 Categories)", 2)
    add_table(doc, ["Category", "Services", "Responsibility"], [
        ["Evidence", "EvidenceDiscoveryService, EvidenceCritiqueService, EvidenceNetworkService",
         "Multi-source search, AI critique, citation network analysis"],
        ["Analysis", "StatisticalAnalysisService, BayesianService, InterimAnalysisService, MissingDataService",
         "Cox PH, IPTW, Bayesian posteriors, group-sequential boundaries, MI/MMRM"],
        ["Data", "CdiscService, SdtmService, AdamService, DataValidationService",
         "CDISC dataset generation, cross-domain validation, metadata export"],
        ["Regulatory", "SarPipelineService, CsrService, SapService, AdrgService, DefineXmlService, EctdService",
         "Document generation, eCTD assembly, validation"],
        ["Review", "CollaborativeReviewService, ReviewConflictService",
         "Multi-reviewer workflows, conflict detection and resolution"],
        ["Search", "SemanticSearchService, HybridSearchService, SemanticScholarService",
         "Vector similarity, BM25 + semantic hybrid, external API integration"],
        ["Platform", "AuthService, AuditService, TaskService, NotificationService, FileStorageService",
         "Authentication, audit logging, background tasks, email, S3 storage"],
    ])

    add_heading(doc, "3.3 Core Layer (10 Modules)", 2)
    add_table(doc, ["Module", "Purpose"], [
        ["models.py", "SQLAlchemy ORM models (24 tables across 5 domains)"],
        ["schemas.py", "Pydantic request/response schemas with validation"],
        ["database.py", "Database engine, session factory, connection pooling config"],
        ["security.py", "JWT creation/validation, password hashing (bcrypt), RBAC decorators"],
        ["config.py", "Environment-based configuration (DATABASE_URL, REDIS_URL, S3, Sentry DSN)"],
        ["exceptions.py", "Custom exception hierarchy with HTTP status mapping"],
        ["middleware.py", "CORS, request logging, rate limiting, error handling"],
        ["cache.py", "Redis wrapper with TTL-based caching and invalidation"],
        ["tasks.py", "Background task infrastructure (queue, worker, progress tracking)"],
        ["utils.py", "Shared utilities: date parsing, hashing, file I/O, statistical helpers"],
    ])

    # ═══ Data Flow ═══
    add_heading(doc, "4. Data Flow: 9-Stage Pipeline", 1)
    add_code_block(doc,
        "Stage 1          Stage 2          Stage 3          Stage 4\n"
        "+----------+     +----------+     +----------+     +----------+\n"
        "| Ingest   | --> | Dedupe   | --> | Score    | --> | Critique |\n"
        "| (5 APIs) |     | (DOI +   |     | (Relev- |     | (AI reg- |\n"
        "|          |     |  title)  |     |  ance)   |     |  ulatory)|\n"
        "+----------+     +----------+     +----------+     +----------+\n"
        "                                                        |\n"
        "                                                        v\n"
        "Stage 5          Stage 6          Stage 7          Stage 8\n"
        "+----------+     +----------+     +----------+     +----------+\n"
        "| Review   | --> | Analyze  | --> | Generate | --> | Validate |\n"
        "| (Human   |     | (Stats,  |     | (TFLs,   |     | (CDISC,  |\n"
        "|  decide) |     |  bias)   |     |  docs)   |     |  eCTD)   |\n"
        "+----------+     +----------+     +----------+     +----------+\n"
        "                                                        |\n"
        "                                                        v\n"
        "                                                   Stage 9\n"
        "                                                   +----------+\n"
        "                                                   | Package  |\n"
        "                                                   | (eCTD    |\n"
        "                                                   |  Module5)|\n"
        "                                                   +----------+")

    add_table(doc, ["Stage", "Name", "Input", "Output", "Key Services"], [
        ["1", "Ingest", "Search query + config", "Raw evidence records", "EvidenceDiscoveryService"],
        ["2", "Dedupe", "Raw records", "Unique evidence set", "DOI matching, title cosine similarity"],
        ["3", "Score", "Unique evidence", "Scored evidence (0-100)", "AI relevance scoring"],
        ["4", "Critique", "Scored evidence", "Regulatory critique annotations", "EvidenceCritiqueService"],
        ["5", "Review", "Critiqued evidence", "Include/exclude decisions", "CollaborativeReviewService"],
        ["6", "Analyze", "Included evidence + data", "Statistical results, bias assessment", "StatisticalAnalysisService"],
        ["7", "Generate", "Analysis results", "TFLs, CSR, SAP, ADRG, Define-XML", "CsrService, TflService"],
        ["8", "Validate", "Generated artifacts", "Validation report", "DataValidationService"],
        ["9", "Package", "Validated artifacts", "eCTD Module 5 ZIP", "EctdService"],
    ])

    # ═══ ERD ═══
    add_heading(doc, "5. Entity-Relationship Diagram", 1)
    add_para(doc, "24 tables organized into 5 domains. Foreign key relationships shown below.")

    add_heading(doc, "5.1 Identity Domain", 2)
    add_code_block(doc,
        "+------------------+       +------------------+\n"
        "|   organizations  |       |      users       |\n"
        "+------------------+       +------------------+\n"
        "| PK  id           |<------| FK  org_id       |\n"
        "|     name          |       | PK  id           |\n"
        "|     slug          |       |     email        |\n"
        "|     jurisdiction  |       |     hashed_pw    |\n"
        "|     created_at    |       |     role         |\n"
        "+------------------+       |     is_active     |\n"
        "                           |     created_at    |\n"
        "                           +------------------+\n"
        "                                   |\n"
        "                           +-------+--------+\n"
        "                           |  refresh_tokens |\n"
        "                           +----------------+\n"
        "                           | PK  id          |\n"
        "                           | FK  user_id     |\n"
        "                           |     token_hash  |\n"
        "                           |     expires_at  |\n"
        "                           |     revoked     |\n"
        "                           +----------------+")

    add_heading(doc, "5.2 Project Domain", 2)
    add_code_block(doc,
        "+------------------+       +------------------+       +------------------+\n"
        "|     projects     |       |    evidence      |       |    decisions     |\n"
        "+------------------+       +------------------+       +------------------+\n"
        "| PK  id           |<------| FK  project_id   |<------| FK  evidence_id  |\n"
        "| FK  owner_id     |       | PK  id           |       | PK  id           |\n"
        "| FK  org_id       |       |     source       |       | FK  reviewer_id  |\n"
        "|     title         |       |     title        |       |     decision     |\n"
        "|     status        |       |     doi          |       |     rationale    |\n"
        "|     therapeutic   |       |     abstract     |       |     signature    |\n"
        "|     proc_config   |       |     relevance    |       |     signed_at    |\n"
        "|     created_at    |       |     metadata     |       +------------------+\n"
        "+------------------+       |     created_at    |\n"
        "        |                  +------------------+\n"
        "        |\n"
        "+-------+----------+       +------------------+\n"
        "|    artifacts     |       |  audit_events    |\n"
        "+------------------+       +------------------+\n"
        "| PK  id           |       | PK  id           |\n"
        "| FK  project_id   |       | FK  project_id   |\n"
        "|     artifact_type|       | FK  user_id      |\n"
        "|     file_path    |       |     event_type   |\n"
        "|     metadata     |       |     payload      |\n"
        "|     created_at   |       |     payload_hash |\n"
        "+------------------+       |     created_at   |\n"
        "                           +------------------+")

    add_heading(doc, "5.3 Review Domain", 2)
    add_code_block(doc,
        "+------------------+       +------------------+       +------------------+\n"
        "| review_workflows |       | review_assigns   |       | review_comments  |\n"
        "+------------------+       +------------------+       +------------------+\n"
        "| PK  id           |<------| FK  workflow_id  |       | PK  id           |\n"
        "| FK  project_id   |       | PK  id           |       | FK  evidence_id  |\n"
        "|     status        |       | FK  reviewer_id  |       | FK  user_id      |\n"
        "|     created_at    |       | FK  evidence_id  |       |     body         |\n"
        "+------------------+       |     status        |       |     parent_id    |\n"
        "                           +------------------+       |     created_at   |\n"
        "                                                       +------------------+")

    add_heading(doc, "5.4 Analysis Domain", 2)
    add_code_block(doc,
        "+------------------+       +------------------+       +------------------+\n"
        "| adam_datasets     |       | sdtm_domains     |       |  tfl_outputs     |\n"
        "+------------------+       +------------------+       +------------------+\n"
        "| PK  id           |       | PK  id           |       | PK  id           |\n"
        "| FK  project_id   |       | FK  project_id   |       | FK  project_id   |\n"
        "|     dataset_type |       |     domain_code  |       |     tfl_type     |\n"
        "|     data_json    |       |     data_json    |       |     file_path    |\n"
        "|     metadata     |       |     validated    |       |     metadata     |\n"
        "|     created_at   |       |     created_at   |       |     created_at   |\n"
        "+------------------+       +------------------+       +------------------+")

    add_heading(doc, "5.5 System Domain", 2)
    add_code_block(doc,
        "+------------------+       +------------------+       +------------------+\n"
        "| background_tasks |       | saved_searches   |       |  fed_nodes       |\n"
        "+------------------+       +------------------+       +------------------+\n"
        "| PK  id           |       | PK  id           |       | PK  id           |\n"
        "| FK  user_id      |       | FK  user_id      |       |     name         |\n"
        "|     task_type    |       |     query        |       |     url          |\n"
        "|     status       |       |     filters      |       |     status       |\n"
        "|     progress     |       |     created_at   |       |     last_seen    |\n"
        "|     result       |       +------------------+       +------------------+\n"
        "|     created_at   |\n"
        "+------------------+")

    add_para(doc, "Full table list (24 tables):", bold=True)
    add_table(doc, ["Domain", "Tables", "Count"], [
        ["Identity", "organizations, users, refresh_tokens, password_reset_tokens", "4"],
        ["Project", "projects, evidence, decisions, artifacts, audit_events", "5"],
        ["Review", "review_workflows, review_assignments, review_comments, review_conflicts", "4"],
        ["Analysis", "adam_datasets, sdtm_domains, tfl_outputs, analysis_results, bayesian_results, interim_results", "6"],
        ["System", "background_tasks, saved_searches, federated_nodes, evidence_patterns, notification_log", "5"],
    ])

    # ═══ Enums ═══
    add_heading(doc, "6. Enum Reference", 1)
    add_table(doc, ["Enum", "Values", "Used In"], [
        ["ProjectStatus", "draft, processing, review, completed, archived", "projects.status"],
        ["EvidenceSource", "pubmed, clinicaltrials, openalex, semantic_scholar, biogpt", "evidence.source"],
        ["ReviewDecision", "include, exclude, uncertain, escalate", "decisions.decision"],
        ["UserRole", "admin, biostatistician, reviewer, viewer", "users.role"],
        ["ArtifactType", "sar, csr_synopsis, csr_section_11, csr_section_12, csr_appendix_16, sap, adrg, define_xml, ectd, tfl_package", "artifacts.artifact_type"],
    ])

    # ═══ Dependency Graph ═══
    add_heading(doc, "7. Dependency Graph", 1)
    add_para(doc, "Workflow steps have prerequisite dependencies. A step cannot execute until its "
             "dependencies are satisfied.")
    add_code_block(doc,
        "Step 1: Study Definition\n"
        "  |\n"
        "  +--> Step 2: Causal Framework\n"
        "  |      |\n"
        "  |      +--> Step 5: Comparability & Balance\n"
        "  |             |\n"
        "  |             +--> Step 6: Bias & Sensitivity\n"
        "  |                    |\n"
        "  |                    +--> Step 7: Effect Estimation\n"
        "  |                           |\n"
        "  |                           +--> Step 8: Reproducibility\n"
        "  |                           |      |\n"
        "  |                           |      +--> Step 10: Regulatory Output\n"
        "  |                           |\n"
        "  |                           +--> Step 9: Audit Trail (continuous)\n"
        "  |\n"
        "  +--> Step 3: Data Provenance\n"
        "         |\n"
        "         +--> Step 4: Cohort Construction\n"
        "                |\n"
        "                +--> Step 5 (merges with Causal Framework)")

    add_table(doc, ["Step", "Depends On", "Produces"], [
        ["1. Study Definition", "None", "Protocol hash, study parameters"],
        ["2. Causal Framework", "Step 1", "DAG, adjustment set, covariate list"],
        ["3. Data Provenance", "Step 1", "Evidence records with provenance metadata"],
        ["4. Cohort Construction", "Step 3", "Filtered cohort, attrition funnel"],
        ["5. Comparability", "Steps 2, 4", "PS model, IPTW weights, SMD values"],
        ["6. Bias & Sensitivity", "Step 5", "Bias ratings, E-values, sensitivity results"],
        ["7. Effect Estimation", "Step 6", "Primary analysis, subgroups, forest plot data"],
        ["8. Reproducibility", "Step 7", "Environment manifest, package versions"],
        ["9. Audit Trail", "All steps (continuous)", "Immutable event log"],
        ["10. Regulatory Output", "Steps 7, 8", "TFLs, CSR, SAP, ADRG, Define-XML, eCTD"],
    ])

    doc.save(filepath)
    sz = os.path.getsize(filepath)
    print(f"  -> {filepath} ({sz:,} bytes)")


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("Afarensis Enterprise Documentation Generator (Part 2/2)")
    print("=" * 60)
    gen_howto()
    gen_annotated_report()
    gen_architecture()
    print("\nPart 2 complete (3/6 documents generated).")
