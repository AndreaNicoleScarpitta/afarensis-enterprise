#!/usr/bin/env python3
"""Generate Afarensis Enterprise docs 1-3: API PDF, Reference Guide, Tutorial."""

import sys, os, json
from datetime import date

sys.path.insert(0, os.path.join(os.path.dirname(__file__), "..", "backend"))

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

from reportlab.lib.pagesizes import letter
from reportlab.lib.colors import HexColor, white, black
from reportlab.lib.units import inch
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table,
                                 TableStyle, PageBreak, KeepTogether)
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT

DOCS_DIR = os.path.dirname(os.path.abspath(__file__))
BACKEND_DIR = os.path.join(os.path.dirname(DOCS_DIR), "backend")
TODAY = date.today().strftime("%B %d, %Y")
NAVY = RGBColor(0x1E, 0x3A, 0x5F)
BLUE = RGBColor(0x25, 0x63, 0xEB)

# ── Load endpoints ──────────────────────────────────────────────────────────
with open(os.path.join(BACKEND_DIR, "_api_docs.json")) as f:
    api_data = json.load(f)
ENDPOINTS = api_data["endpoints"]

# ── Helpers ─────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def docx_title_page(doc, title, subtitle="", version="v2.1"):
    for _ in range(6):
        doc.add_paragraph("")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("AFARENSIS ENTERPRISE")
    run.font.size = Pt(28)
    run.font.color.rgb = NAVY
    run.bold = True
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(title)
    run2.font.size = Pt(20)
    run2.font.color.rgb = NAVY
    run2.font.name = "Calibri"

    if subtitle:
        p3 = doc.add_paragraph()
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run3 = p3.add_run(subtitle)
        run3.font.size = Pt(14)
        run3.font.color.rgb = BLUE
        run3.font.name = "Calibri"

    doc.add_paragraph("")
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = p4.add_run(f"Version {version}  |  {TODAY}")
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run4.font.name = "Calibri"

    p5 = doc.add_paragraph()
    p5.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run5 = p5.add_run("Synthetic Ascension Pty Ltd")
    run5.font.size = Pt(11)
    run5.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    run5.font.name = "Calibri"

    doc.add_page_break()

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Calibri"
    return h

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(11)
    run.bold = bold
    run.italic = italic
    return p

def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)
    return p

def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, "1E3A5F")
    # Data rows
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = "Calibri"
                    run.font.size = Pt(9)
            if ri % 2 == 1:
                set_cell_shading(cell, "F0F4FA")
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph("")
    return table


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENT 1: API Reference PDF (reportlab)
# ═════════════════════════════════════════════════════════════════════════════

def categorize_endpoints(endpoints):
    categories = {}
    for ep in endpoints:
        path = ep["path"]
        if "/auth/" in path:
            cat = "Auth"
        elif "/tasks" in path:
            cat = "Tasks"
        elif "/projects/" in path and "/study/" in path and "/tfl/" in path:
            cat = "TFL"
        elif "/projects/" in path and "/study/" in path and "/bayesian/" in path:
            cat = "Bayesian"
        elif "/projects/" in path and "/study/" in path and "/interim/" in path:
            cat = "Interim"
        elif "/projects/" in path and "/study/" in path and "/missing-data/" in path:
            cat = "Statistics"
        elif "/projects/" in path and "/study/" in path and "/regulatory" in path:
            cat = "Regulatory"
        elif "/projects/" in path and "/study/" in path:
            cat = "Study Config"
        elif "/projects/" in path and "/adam/" in path:
            cat = "CDISC"
        elif "/projects/" in path and "/sdtm/" in path:
            cat = "SDTM"
        elif "/projects/" in path and "/submission/" in path:
            cat = "Regulatory"
        elif "/projects/" in path and "/evidence" in path:
            cat = "Evidence"
        elif "/projects/" in path and "/bias" in path:
            cat = "Evidence"
        elif "/review/" in path:
            cat = "Review"
        elif "/search/" in path:
            cat = "Search"
        elif "/statistics/" in path:
            cat = "Statistics"
        elif "/sar-pipeline/" in path:
            cat = "Regulatory"
        elif "/program/" in path:
            cat = "Program"
        elif "/org/" in path:
            cat = "Org Management"
        elif "/system/" in path or "/health" in path:
            cat = "System"
        elif "/biogpt/" in path:
            cat = "BioGPT"
        elif "/users/" in path or "/user/" in path:
            cat = "Org Management"
        elif "/federated/" in path:
            cat = "System"
        elif "/evidence-patterns" in path:
            cat = "Evidence"
        elif "/analytics/" in path:
            cat = "System"
        elif "/audit/" in path:
            cat = "System"
        elif "/data/" in path:
            cat = "System"
        elif "/projects" in path:
            cat = "Projects"
        else:
            cat = "System"
        categories.setdefault(cat, []).append(ep)
    return categories

def gen_api_pdf():
    filepath = os.path.join(DOCS_DIR, "Afarensis_API_Reference_v2.1.pdf")
    print(f"  Generating API Reference PDF...")
    doc = SimpleDocTemplate(filepath, pagesize=letter,
                            topMargin=0.75*inch, bottomMargin=0.75*inch,
                            leftMargin=0.75*inch, rightMargin=0.75*inch)

    styles = getSampleStyleSheet()
    navy = HexColor("#1E3A5F")
    blue = HexColor("#2563EB")

    title_style = ParagraphStyle("Title2", parent=styles["Title"],
                                  textColor=navy, fontSize=24, spaceAfter=6,
                                  fontName="Helvetica-Bold", alignment=TA_CENTER)
    subtitle_style = ParagraphStyle("Sub", parent=styles["Normal"],
                                     textColor=HexColor("#666666"), fontSize=11,
                                     alignment=TA_CENTER, spaceAfter=20)
    cat_style = ParagraphStyle("Cat", parent=styles["Heading1"],
                                textColor=navy, fontSize=16, spaceAfter=8,
                                spaceBefore=16, fontName="Helvetica-Bold")
    normal = ParagraphStyle("Norm", parent=styles["Normal"], fontSize=9,
                             fontName="Helvetica", spaceAfter=4)
    small = ParagraphStyle("Small", parent=styles["Normal"], fontSize=8,
                            fontName="Helvetica", textColor=HexColor("#444444"))
    header_style = ParagraphStyle("Hdr", parent=styles["Normal"], fontSize=9,
                                   fontName="Helvetica-Bold", textColor=white)

    elements = []

    # Title page
    elements.append(Spacer(1, 2*inch))
    elements.append(Paragraph("AFARENSIS ENTERPRISE", title_style))
    elements.append(Paragraph("API Reference v2.1", ParagraphStyle(
        "T2", parent=title_style, fontSize=18, spaceAfter=20)))
    elements.append(Paragraph(f"146 Endpoints  |  18 Categories  |  {TODAY}", subtitle_style))
    elements.append(Paragraph("Synthetic Ascension Pty Ltd", subtitle_style))
    elements.append(PageBreak())

    # TOC
    elements.append(Paragraph("Table of Contents", cat_style))
    categories = categorize_endpoints(ENDPOINTS)
    cat_order = ["Auth", "Projects", "Evidence", "Review", "Study Config", "CDISC",
                 "TFL", "Regulatory", "Search", "Statistics", "Bayesian", "Interim",
                 "SDTM", "Program", "System", "BioGPT", "Tasks", "Org Management"]
    for cat in cat_order:
        eps = categories.get(cat, [])
        if eps:
            elements.append(Paragraph(f"  {cat} ({len(eps)} endpoints)", normal))
    elements.append(PageBreak())

    # RBAC Roles
    elements.append(Paragraph("RBAC Roles", cat_style))
    role_data = [
        ["admin", "Full access", "User mgmt, system config, all projects"],
        ["biostatistician", "Analysis access", "Create/run studies, generate TFLs, stats"],
        ["reviewer", "Review access", "Review evidence, submit decisions, comment"],
        ["viewer", "Read-only", "View projects, download artifacts"],
    ]
    role_table = Table(
        [["Role", "Access Level", "Permissions"]] + role_data,
        colWidths=[1.5*inch, 1.5*inch, 4*inch]
    )
    role_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), navy),
        ("TEXTCOLOR", (0,0), (-1,0), white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("GRID", (0,0), (-1,-1), 0.5, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [white, HexColor("#F0F4FA")]),
        ("VALIGN", (0,0), (-1,-1), "TOP"),
    ]))
    elements.append(role_table)
    elements.append(Spacer(1, 12))

    # Default accounts
    elements.append(Paragraph("Default Accounts", cat_style))
    acct_data = [
        ["admin@afarensis.io", "admin", "Admin123!", "Full system access"],
        ["bio@afarensis.io", "biostatistician", "Bio123!", "Analysis workspace"],
        ["reviewer@afarensis.io", "reviewer", "Review123!", "Review queue"],
    ]
    acct_table = Table(
        [["Email", "Role", "Password", "Access"]] + acct_data,
        colWidths=[2*inch, 1.5*inch, 1.3*inch, 2.2*inch]
    )
    acct_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), navy),
        ("TEXTCOLOR", (0,0), (-1,0), white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("GRID", (0,0), (-1,-1), 0.5, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [white, HexColor("#F0F4FA")]),
    ]))
    elements.append(acct_table)
    elements.append(Spacer(1, 12))

    # Error format
    elements.append(Paragraph("Error Response Format", cat_style))
    elements.append(Paragraph('All errors return: {"detail": "message"} with appropriate HTTP status codes.', normal))
    err_data = [
        ["400", "Bad Request", "Invalid input or missing required fields"],
        ["401", "Unauthorized", "Missing or expired JWT token"],
        ["403", "Forbidden", "Insufficient role permissions"],
        ["404", "Not Found", "Resource does not exist"],
        ["409", "Conflict", "Duplicate or state conflict"],
        ["422", "Validation Error", "Request body fails schema validation"],
        ["429", "Rate Limited", "Too many requests"],
        ["500", "Server Error", "Internal error (logged to Sentry)"],
    ]
    err_table = Table(
        [["Code", "Status", "Description"]] + err_data,
        colWidths=[1*inch, 1.5*inch, 4.5*inch]
    )
    err_table.setStyle(TableStyle([
        ("BACKGROUND", (0,0), (-1,0), navy),
        ("TEXTCOLOR", (0,0), (-1,0), white),
        ("FONTNAME", (0,0), (-1,0), "Helvetica-Bold"),
        ("FONTSIZE", (0,0), (-1,-1), 9),
        ("GRID", (0,0), (-1,-1), 0.5, HexColor("#CCCCCC")),
        ("ROWBACKGROUNDS", (0,1), (-1,-1), [white, HexColor("#F0F4FA")]),
    ]))
    elements.append(err_table)
    elements.append(PageBreak())

    # Method colors
    method_colors = {
        "GET": HexColor("#16A34A"),
        "POST": HexColor("#2563EB"),
        "PUT": HexColor("#EA580C"),
        "DELETE": HexColor("#DC2626"),
        "PATCH": HexColor("#9333EA"),
    }

    # Endpoints by category
    for cat in cat_order:
        eps = categories.get(cat, [])
        if not eps:
            continue
        elements.append(Paragraph(f"{cat} ({len(eps)} endpoints)", cat_style))

        rows = []
        for ep in eps:
            method = ep["method"]
            mc = method_colors.get(method, black)
            method_p = Paragraph(f'<font color="{mc.hexval()}">{method}</font>',
                                  ParagraphStyle("M", fontSize=9, fontName="Helvetica-Bold"))
            path_p = Paragraph(f'<font name="Courier" size="8">{ep["path"]}</font>', normal)
            doc_text = (ep.get("doc") or "").split("\n")[0][:120]
            desc_p = Paragraph(doc_text, small)
            auth = "JWT" if "/auth/login" not in ep["path"] and "/health" not in ep["path"] else "None"
            auth_p = Paragraph(auth, small)
            rows.append([method_p, path_p, desc_p, auth_p])

        header_row = [
            Paragraph("Method", header_style),
            Paragraph("Path", header_style),
            Paragraph("Description", header_style),
            Paragraph("Auth", header_style),
        ]
        t = Table([header_row] + rows, colWidths=[0.7*inch, 2.8*inch, 3*inch, 0.5*inch])
        t.setStyle(TableStyle([
            ("BACKGROUND", (0,0), (-1,0), navy),
            ("TEXTCOLOR", (0,0), (-1,0), white),
            ("GRID", (0,0), (-1,-1), 0.5, HexColor("#DDDDDD")),
            ("ROWBACKGROUNDS", (0,1), (-1,-1), [white, HexColor("#F8FAFC")]),
            ("VALIGN", (0,0), (-1,-1), "TOP"),
            ("TOPPADDING", (0,0), (-1,-1), 3),
            ("BOTTOMPADDING", (0,0), (-1,-1), 3),
        ]))
        elements.append(t)
        elements.append(Spacer(1, 8))

    doc.build(elements)
    sz = os.path.getsize(filepath)
    print(f"  -> {filepath} ({sz:,} bytes)")


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENT 2: Reference Guide (DOCX)
# ═════════════════════════════════════════════════════════════════════════════

def gen_reference_guide():
    filepath = os.path.join(DOCS_DIR, "Afarensis_Reference_Guide.docx")
    print(f"  Generating Reference Guide...")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    docx_title_page(doc, "Reference Guide", "Product Manager Edition")

    # ── Section 1: Welcome ──
    add_heading(doc, "1. Welcome to Afarensis", 1)
    add_para(doc, "Afarensis Enterprise is a regulatory-grade evidence synthesis and clinical study "
             "reporting platform. It automates the journey from literature discovery through "
             "statistical analysis to eCTD submission packaging, enabling biostatisticians, "
             "regulatory writers, and HEOR professionals to produce ICH E9(R1)-compliant "
             "deliverables in a fraction of the time.")
    add_para(doc, "Who is it for?", bold=True)
    add_bullet(doc, "Biostatisticians designing and executing causal inference studies")
    add_bullet(doc, "Regulatory medical writers assembling CSR, ADRG, and Define-XML")
    add_bullet(doc, "HEOR professionals conducting comparative effectiveness research")
    add_bullet(doc, "Clinical data managers producing CDISC ADaM and SDTM datasets")
    add_bullet(doc, "Quality assurance teams auditing study conduct for 21 CFR Part 11")
    add_para(doc, "What problem does it solve?", bold=True)
    add_para(doc, "Traditional evidence synthesis requires manual searches across multiple databases, "
             "copying results into spreadsheets, running SAS programs offline, and manually assembling "
             "regulatory documents. Afarensis unifies these steps into a single, auditable, reproducible "
             "workflow with built-in regulatory compliance.")

    # ── Section 2: The Interface ──
    add_heading(doc, "2. The Interface", 1)
    add_heading(doc, "2.1 Dashboard Layout", 2)
    add_para(doc, "The main dashboard displays project cards in a responsive grid. Each card shows the "
             "project name, therapeutic area, current status badge, evidence count, and progress bar "
             "indicating completion through the 10-step workflow.")
    add_heading(doc, "2.2 Sidebar Navigation", 2)
    add_para(doc, "The left sidebar contains the 10-step workflow navigator. Each step shows an icon, "
             "label, and status indicator (completed checkmark, in-progress spinner, or locked padlock). "
             "Steps unlock sequentially as prerequisites are met.")
    add_heading(doc, "2.3 Top Bar", 2)
    add_para(doc, "The top bar displays the Afarensis logo, current project name, user avatar with role badge, "
             "notification bell, and a global search field. The search supports semantic queries across all "
             "evidence records.")
    add_heading(doc, "2.4 Project Cards", 2)
    add_para(doc, "Each project card surfaces: project title, therapeutic area tag, status badge (Draft / "
             "Processing / Review / Completed / Archived), evidence record count, last-modified timestamp, "
             "and the assigned lead biostatistician.")

    # ── Section 3: Visual Language ──
    add_heading(doc, "3. Visual Language", 1)
    add_para(doc, "Afarensis uses a consistent color system across the entire interface to communicate "
             "meaning at a glance.")
    colors = [
        ["Navy (#1E3A5F)", "Headers, primary brand, navigation backgrounds, table headers"],
        ["Blue (#2563EB)", "Interactive elements: links, buttons, active sidebar items, focused inputs"],
        ["Green badges", "Completed items, passing validations, successful operations, balanced covariates"],
        ["Orange / Amber badges", "Warnings, in-progress items, approaching thresholds (e.g., SMD near 0.10)"],
        ["Red badges", "Errors, failed validations, critical alerts, unbalanced covariates"],
        ["Gray", "Disabled elements, secondary text, locked steps, placeholder text"],
    ]
    add_table(doc, ["Color", "Usage"], colors)

    # ── Section 4: Project Lifecycle ──
    add_heading(doc, "4. Project Lifecycle", 1)
    add_para(doc, "Every project moves through five states. Transitions are triggered by specific "
             "actions and are logged in the audit trail.")
    states = [
        ["Draft", "Project created; study definition in progress", "Complete study definition and lock protocol"],
        ["Processing", "Evidence discovery, analysis, and generation running", "All 10 steps completed and reviewed"],
        ["Review", "Reviewer(s) assessing evidence decisions and outputs", "All review decisions submitted with e-signatures"],
        ["Completed", "All outputs generated and approved", "Admin archives or exports eCTD"],
        ["Archived", "Read-only historical record", "Admin unarchives (returns to Completed)"],
    ]
    add_table(doc, ["State", "Description", "Transition Trigger"], states)

    # ── Section 5: The 10-Step Workflow ──
    add_heading(doc, "5. The 10-Step Workflow", 1)

    steps = [
        ("Study Definition", [
            "Configure the fundamental parameters of your study: title, therapeutic area, indication, "
            "study design (RCT, cohort, case-control), primary endpoint, treatment and comparator arms.",
            "Form fields include: Study ID (auto-generated), Study Title, Indication (ICD-10 coded), "
            "Phase (I-IV or observational), Design Type, Primary Endpoint (time-to-event, binary, continuous), "
            "Treatment Arm description, Comparator Arm description, Target Sample Size, and Alpha Level.",
            "Once locked via the Lock Protocol button, the study definition becomes immutable and an audit "
            "log entry records the locking user, timestamp, and protocol hash."
        ]),
        ("Causal Framework", [
            "The DAG (Directed Acyclic Graph) editor allows you to define causal relationships between "
            "variables. Select covariates from a categorized list (demographics, clinical, laboratory, "
            "concomitant medications) and draw directed edges to specify assumed causal paths.",
            "The system identifies confounders, mediators, and colliders, then recommends an adjustment set. "
            "Covariates in the adjustment set are automatically carried forward to the balance assessment step."
        ]),
        ("Data Provenance", [
            "Evidence sources are searched in parallel. Each source displays a colored badge indicating "
            "its origin: PubMed (blue), ClinicalTrials.gov (green), OpenAlex (orange), Semantic Scholar "
            "(purple), and BioGPT (pink, for AI-generated summaries).",
            "Records are deduplicated by DOI and title similarity. Each record receives a relevance score "
            "(0-100) computed by the AI critique engine. Provenance metadata includes retrieval timestamp, "
            "query parameters, and source API response identifiers."
        ]),
        ("Cohort Construction", [
            "Define inclusion and exclusion criteria as structured rules. Each criterion specifies a "
            "variable, operator (equals, greater than, contains, between), and threshold value.",
            "The attrition funnel visualization shows how many subjects remain after each criterion is "
            "applied sequentially. Green bars indicate criteria retaining > 80% of subjects; amber bars "
            "indicate 50-80% retention; red bars indicate < 50% retention."
        ]),
        ("Comparability & Balance", [
            "The Love plot displays the Standardized Mean Difference (SMD) for each covariate before "
            "and after propensity score adjustment (IPTW). Covariates with SMD < 0.10 appear as green "
            "dots (balanced); those with SMD between 0.10 and 0.25 appear as amber dots; and those with "
            "SMD > 0.25 appear as red dots (critically unbalanced).",
            "The 0.10 threshold is the conventional standard recommended by Austin (2011) for declaring "
            "covariate balance. The system flags any covariate exceeding this threshold and suggests "
            "remedial actions: additional covariates, different PS model, or trimming extreme weights."
        ]),
        ("Bias & Sensitivity", [
            "Five bias types are assessed: selection bias, information bias, confounding, attrition bias, "
            "and reporting bias. Each receives a severity rating (low, moderate, high, critical).",
            "The E-value quantifies the minimum strength of association that an unmeasured confounder "
            "would need with both treatment and outcome to explain away the observed effect. Higher "
            "E-values indicate more robust results.",
            "The sensitivity catalog offers: Rosenbaum bounds, Manski bounds, tipping-point analysis, "
            "probabilistic bias analysis, and negative control outcome analysis."
        ]),
        ("Effect Estimation", [
            "The forest plot displays the primary analysis result (e.g., Cox PH hazard ratio) as a "
            "diamond, with subgroup analyses and sensitivity analyses as individual squares with "
            "confidence intervals.",
            "The Show Your Work drawer contains six tabs: (1) Model Specification, showing the exact "
            "regression formula; (2) Diagnostics, with residual plots; (3) Assumptions, listing "
            "checked model assumptions; (4) Code, the reproducible Python/R code; (5) Interpretation, "
            "a plain-English summary; (6) Audit, the timestamp and user who ran the analysis.",
            "Multiplicity adjustment methods include: Bonferroni, Holm, Hochberg, Benjamini-Hochberg "
            "(FDR), and fixed-sequence testing."
        ]),
        ("Reproducibility", [
            "The environment manifest records: Python version, all package versions (pip freeze), "
            "operating system, CPU architecture, random seed, and database schema hash.",
            "Package versions are pinned at analysis execution time. Any subsequent run with different "
            "versions triggers a warning in the audit trail."
        ]),
        ("Audit Trail", [
            "Every significant action is logged: study definition changes, evidence decisions, analysis "
            "runs, document generation, user logins, role changes, and artifact downloads.",
            "Each audit entry contains: event type, timestamp (UTC), user ID, user role, IP address, "
            "previous value (for changes), new value, and a SHA-256 hash of the event payload.",
            "Retention policy: 7 years, consistent with ICH E6(R2) GCP requirements and 21 CFR Part 11 "
            "electronic records regulations. Audit logs cannot be modified or deleted."
        ]),
        ("Regulatory Output", [
            "Tables, Figures, and Listings (TFLs): Demographics table (Table 14.1.1), Adverse Events "
            "table (Table 14.3.1), Kaplan-Meier curves (Figure 14.2.1), Forest plot (Figure 14.2.2), "
            "and Love plot (Figure 14.1.1).",
            "Regulatory documents: Statistical Analysis Plan (SAP), Clinical Study Report (CSR) with "
            "Synopsis, Section 11 (Efficacy), Section 12 (Safety), and Appendix 16.",
            "Submission artifacts: Analysis Data Reviewer's Guide (ADRG), Define-XML 2.1, eCTD Module 5 "
            "package with Study Tagging File, SDTM domains (DM, AE, LB, VS, EX, DS), and ADaM datasets "
            "(ADSL, ADAE, ADTTE)."
        ]),
    ]

    for i, (name, paragraphs) in enumerate(steps, 1):
        add_heading(doc, f"5.{i} Step {i}: {name}", 2)
        for text in paragraphs:
            add_para(doc, text)

    # ── Section 6: Evidence Source Integration ──
    add_heading(doc, "6. Evidence Source Integration", 1)
    sources = [
        ["PubMed", "Blue", "NCBI E-utilities API", "Peer-reviewed biomedical literature",
         "PMID, title, abstract, MeSH terms, publication date, authors, journal"],
        ["ClinicalTrials.gov", "Green", "CT.gov v2 API", "Registered clinical trials",
         "NCT ID, status, conditions, interventions, outcome measures, enrollment"],
        ["OpenAlex", "Orange", "OpenAlex REST API", "Open-access scholarly works",
         "DOI, title, abstract, cited-by count, concepts, host venue"],
        ["Semantic Scholar", "Purple", "S2 Academic Graph API", "Academic papers with citation graphs",
         "S2 Paper ID, title, abstract, citation count, influential citations, TLDR"],
        ["BioGPT", "Pink", "Local model inference", "AI-generated biomedical summaries",
         "Generated text, confidence score, source references, mechanism explanations"],
    ]
    add_table(doc, ["Source", "Badge Color", "Integration", "Content Type", "Fields Extracted"], sources)

    # ── Section 7: Collaboration ──
    add_heading(doc, "7. Collaboration Features", 1)
    add_heading(doc, "7.1 Reviewer Assignment", 2)
    add_para(doc, "Admins and lead biostatisticians assign reviewers to specific evidence records or entire "
             "projects. Reviewers receive email notifications and see assignments in their review queue.")
    add_heading(doc, "7.2 Commenting", 2)
    add_para(doc, "Threaded comments attach to individual evidence records. Comments support @mentions "
             "(notifying the mentioned user), markdown formatting, and file attachments.")
    add_heading(doc, "7.3 Conflict Resolution", 2)
    add_para(doc, "When two reviewers disagree on an evidence decision (include vs. exclude), the system "
             "flags the conflict. A senior reviewer or the lead biostatistician resolves the conflict by "
             "selecting the final decision and providing a resolution rationale.")
    add_heading(doc, "7.4 E-Signatures", 2)
    add_para(doc, "Evidence decisions and regulatory document approvals require cryptographic e-signatures "
             "(21 CFR Part 11 compliant). The signature records: user identity, timestamp, meaning of "
             "signature (e.g., 'I approve this evidence for inclusion'), and a SHA-256 hash of the signed content.")

    # ── Section 8: Administration ──
    add_heading(doc, "8. Administration", 1)
    add_heading(doc, "8.1 User Management", 2)
    add_para(doc, "Admins invite users via email, assign roles (admin, biostatistician, reviewer, viewer), "
             "and can deactivate or reactivate accounts. Deactivated users cannot log in but their audit "
             "history is preserved.")
    add_heading(doc, "8.2 Organization Setup", 2)
    add_para(doc, "Each deployment is scoped to a single organization. The organization record stores: "
             "name, regulatory jurisdiction (FDA, EMA, PMDA, TGA), default therapeutic areas, and "
             "branding configuration.")
    add_heading(doc, "8.3 System Health", 2)
    add_para(doc, "The system health dashboard shows: API uptime, database connection pool status, "
             "Redis cache hit rate, background task queue depth, storage utilization, and error rates "
             "aggregated by endpoint category.")

    doc.save(filepath)
    sz = os.path.getsize(filepath)
    print(f"  -> {filepath} ({sz:,} bytes)")


# ═════════════════════════════════════════════════════════════════════════════
# DOCUMENT 3: Tutorial (DOCX)
# ═════════════════════════════════════════════════════════════════════════════

def gen_tutorial():
    filepath = os.path.join(DOCS_DIR, "Afarensis_Tutorial.docx")
    print(f"  Generating Tutorial...")
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    docx_title_page(doc, "End-to-End Tutorial", "From Login to eCTD Submission")

    # ── 1. Login ──
    add_heading(doc, "1. Login and First-Time Setup", 1)
    add_para(doc, "Step 1: Navigate to your Afarensis instance URL (e.g., https://app.afarensis.io).")
    add_para(doc, "Step 2: Enter your email and password. Default accounts are provided for evaluation:")
    add_table(doc, ["Email", "Role", "Password"], [
        ["admin@afarensis.io", "Admin", "Admin123!"],
        ["bio@afarensis.io", "Biostatistician", "Bio123!"],
        ["reviewer@afarensis.io", "Reviewer", "Review123!"],
    ])
    add_para(doc, "Step 3: On first login, you will see the empty dashboard. The system is ready to create "
             "your first project.")

    # ── 2. Creating a Project ──
    add_heading(doc, "2. Creating a Project", 1)
    add_para(doc, "This tutorial uses a worked example: a comparative effectiveness study of surgical "
             "interventions for cystic hygroma in pediatric patients.")
    add_para(doc, 'Step 1: Click the "+ New Project" button on the dashboard.')
    add_para(doc, "Step 2: Fill in the project creation form:")
    add_bullet(doc, 'Project Title: "Cystic Hygroma Surgical Outcomes"')
    add_bullet(doc, "Therapeutic Area: Pediatric Surgery")
    add_bullet(doc, "Indication: Cystic Hygroma (ICD-10: D18.1)")
    add_bullet(doc, "Description: Comparative effectiveness of excision vs. sclerotherapy for cystic hygroma")
    add_para(doc, 'Step 3: Click "Create Project". You will be redirected to the project workspace with the '
             "10-step workflow sidebar visible.")

    # ── 3. Study Protocol ──
    add_heading(doc, "3. Defining the Study Protocol", 1)
    add_para(doc, "Step 1: Click 'Study Definition' in the sidebar (Step 1).")
    add_para(doc, "Step 2: Complete the study definition form:")
    add_bullet(doc, "Study Design: Retrospective Cohort")
    add_bullet(doc, "Primary Endpoint: Time to recurrence (time-to-event)")
    add_bullet(doc, "Treatment Arm: Complete surgical excision")
    add_bullet(doc, "Comparator Arm: OK-432 sclerotherapy")
    add_bullet(doc, "Target Sample Size: 200")
    add_bullet(doc, "Alpha Level: 0.05 (two-sided)")
    add_para(doc, "Step 3: Define covariates in the Causal Framework tab: age at diagnosis, hygroma stage "
             "(I-IV), location (cervical, axillary, mediastinal), prior treatment history, and sex.")
    add_para(doc, 'Step 4: Click "Lock Protocol" to finalize the study definition. This creates an immutable '
             "audit record.")

    # ── 4. Evidence Discovery ──
    add_heading(doc, "4. Running Evidence Discovery", 1)
    add_para(doc, 'Step 1: Navigate to Step 3 (Data Provenance) and click "Discover Evidence".')
    add_para(doc, "Step 2: The system searches five sources simultaneously. Typical results for cystic hygroma:")
    add_table(doc, ["Source", "Badge", "Expected Results", "Key Findings"], [
        ["PubMed", "Blue", "45-60 articles", "Systematic reviews, case series, retrospective studies"],
        ["ClinicalTrials.gov", "Green", "5-10 trials", "Active and completed trials for OK-432"],
        ["OpenAlex", "Orange", "80-120 works", "Broader academic literature including conference papers"],
        ["Semantic Scholar", "Purple", "60-90 papers", "Citation-connected papers with influence scores"],
        ["BioGPT", "Pink", "5 summaries", "AI-synthesized mechanism and outcome summaries"],
    ])
    add_para(doc, "Step 3: Review the deduplicated evidence list. Each record shows: title, source badge, "
             "relevance score, and a brief excerpt.")

    # ── 5. Comparability ──
    add_heading(doc, "5. Reviewing Comparability Scores", 1)
    add_para(doc, 'Step 1: Navigate to Step 5 (Comparability & Balance) and click "Compute Balance".')
    add_para(doc, "Step 2: The system runs propensity score estimation via logistic regression, computes "
             "Inverse Probability of Treatment Weights (IPTW), and calculates SMD for each covariate.")
    add_para(doc, "Step 3: Interpret the Love plot:")
    add_bullet(doc, "Green dots (SMD < 0.10): Covariate is balanced between treatment arms")
    add_bullet(doc, "Amber dots (0.10 - 0.25): Borderline imbalance, consider additional adjustment")
    add_bullet(doc, "Red dots (SMD > 0.25): Significant imbalance requiring remediation")

    # ── 6. Statistical Analysis ──
    add_heading(doc, "6. Running Statistical Analysis", 1)
    add_para(doc, "Step 1: Navigate to Step 7 (Effect Estimation).")
    add_para(doc, "Step 2: The system runs the primary analysis (Cox Proportional Hazards model for "
             "time-to-recurrence) with IPTW adjustment. Results include:")
    add_bullet(doc, "Hazard Ratio with 95% CI")
    add_bullet(doc, "Log-rank test p-value")
    add_bullet(doc, "Kaplan-Meier survival curves for both arms")
    add_para(doc, 'Step 3: Expand the "Show Your Work" drawer to review model diagnostics, code, '
             "and assumptions. The six tabs provide full transparency for regulatory reviewers.")
    add_para(doc, "Step 4: Review sensitivity analyses: propensity score matching, as-treated analysis, "
             "and tipping-point analysis for missing data.")

    # ── 7. CDISC Datasets ──
    add_heading(doc, "7. Generating CDISC Datasets", 1)
    add_para(doc, 'Step 1: Navigate to the CDISC panel and click "Generate ADaM Datasets".')
    add_para(doc, "Step 2: The system generates three ADaM datasets:")
    add_table(doc, ["Dataset", "Description", "Key Variables"], [
        ["ADSL", "Subject-Level Analysis Dataset", "USUBJID, AGE, SEX, TRTA, SAFFL, ITTFL"],
        ["ADAE", "Adverse Events Analysis Dataset", "AEDECOD, AEBODSYS, AESER, AEREL, AESTDTC"],
        ["ADTTE", "Time-to-Event Analysis Dataset", "PARAMCD, AVAL, CNSR, STARTDT, ADT"],
    ])
    add_para(doc, 'Step 3: Click "Validate" to run CDISC conformance checks. Review any findings '
             "and resolve before proceeding.")

    # ── 8. TFLs ──
    add_heading(doc, "8. Creating Tables, Figures, and Listings", 1)
    add_para(doc, 'Step 1: Navigate to the TFL panel and click "Generate All TFLs".')
    add_para(doc, "Step 2: The system generates:")
    add_bullet(doc, "Table 14.1.1: Demographics and Baseline Characteristics")
    add_bullet(doc, "Table 14.3.1: Adverse Events by System Organ Class")
    add_bullet(doc, "Figure 14.1.1: Covariate Balance Love Plot")
    add_bullet(doc, "Figure 14.2.1: Kaplan-Meier Survival Curves")
    add_bullet(doc, "Figure 14.2.2: Forest Plot of Subgroup Analyses")
    add_para(doc, "Step 3: Review each TFL for accuracy and formatting. TFLs are generated in "
             "publication-ready format with proper titles, footnotes, and source annotations.")

    # ── 9. Regulatory Documents ──
    add_heading(doc, "9. Generating Regulatory Documents", 1)
    add_para(doc, "Step 1: Navigate to Step 10 (Regulatory Output).")
    add_para(doc, "Step 2: Generate each document:")
    add_bullet(doc, "SAP (Statistical Analysis Plan): Click Generate SAP. Review the auto-populated "
               "sections including study objectives, endpoints, analysis populations, and statistical methods.")
    add_bullet(doc, "CSR (Clinical Study Report): Click Generate Full CSR. This produces Synopsis, "
               "Section 11 (Efficacy), Section 12 (Safety), and Appendix 16.")
    add_bullet(doc, "ADRG (Analysis Data Reviewer's Guide): Click Generate ADRG. This documents the "
               "computing environment, data flow, and analysis program descriptions.")
    add_bullet(doc, "Define-XML: Click Generate Define-XML. This produces the metadata file describing "
               "all ADaM variables, codelists, and computational methods.")

    # ── 10. eCTD ──
    add_heading(doc, "10. Assembling the eCTD Package", 1)
    add_para(doc, 'Step 1: Click "Generate eCTD Package" in the submission panel.')
    add_para(doc, "Step 2: The system assembles Module 5 (Clinical Study Reports) with the proper "
             "directory structure:")
    add_bullet(doc, "m5/53-clin-stud-rep/535-rep-analys-data-indiv-patient/ : CSR and appendices")
    add_bullet(doc, "m5/datasets/ : ADaM and SDTM datasets with Define-XML")
    add_bullet(doc, "m5/datasets/analysis/ : ADRG and analysis programs")
    add_para(doc, 'Step 3: Click "Validate eCTD" to run structural validation. The system checks for '
             "required files, naming conventions, and cross-references between documents.")
    add_para(doc, "Step 4: Download the validated package as a ZIP file for submission through the "
             "appropriate regulatory gateway (FDA ESG, EMA CESP, etc.).")

    # ── 11. Troubleshooting ──
    add_heading(doc, "11. Troubleshooting", 1)
    issues = [
        ["Evidence discovery returns 0 results", "Broaden search terms; check that the indication text matches PubMed MeSH terms"],
        ["Propensity score model fails to converge", "Reduce the number of covariates; check for perfect separation; increase regularization"],
        ["SMD exceeds 0.10 after weighting", "Add interaction terms to PS model; consider trimming weights > 10; use stabilized weights"],
        ["TFL generation times out", "Large datasets may need increased timeout; check Redis connection; verify background task worker is running"],
        ["Define-XML validation errors", "Check that all variable labels are <= 40 characters; ensure codelists have unique coded values"],
        ["eCTD validation warnings", "Missing recommended files (e.g., reviewer's guide); filename length exceeds 64 characters"],
        ["Login returns 401", "Check that the account is active; verify password; try forgot-password flow"],
        ["Audit trail shows gaps", "Check server timezone configuration; ensure NTP sync is enabled on the server"],
    ]
    add_table(doc, ["Issue", "Resolution"], issues)

    # ── Appendix ──
    add_heading(doc, "Appendix A: Keyboard Shortcuts", 1)
    shortcuts = [
        ["Ctrl + N", "Create new project"],
        ["Ctrl + S", "Save current form"],
        ["Ctrl + Enter", "Run current analysis"],
        ["Ctrl + Shift + E", "Open evidence discovery"],
        ["Ctrl + Shift + T", "Generate all TFLs"],
        ["Ctrl + /", "Open command palette"],
        ["Esc", "Close modal or drawer"],
        ["Tab / Shift+Tab", "Navigate form fields"],
    ]
    add_table(doc, ["Shortcut", "Action"], shortcuts)

    add_heading(doc, "Appendix B: Account Types", 1)
    acct_types = [
        ["Admin", "Full system access including user management, organization settings, and system health monitoring"],
        ["Biostatistician", "Create and manage projects, run analyses, generate TFLs and regulatory documents"],
        ["Reviewer", "Review and adjudicate evidence, submit e-signed decisions, resolve conflicts"],
        ["Viewer", "Read-only access to projects and artifacts; cannot modify data or run analyses"],
    ]
    add_table(doc, ["Account Type", "Capabilities"], acct_types)

    doc.save(filepath)
    sz = os.path.getsize(filepath)
    print(f"  -> {filepath} ({sz:,} bytes)")


# ═════════════════════════════════════════════════════════════════════════════
# MAIN
# ═════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    print("=" * 60)
    print("Afarensis Enterprise Documentation Generator (Part 1/2)")
    print("=" * 60)
    gen_api_pdf()
    gen_reference_guide()
    gen_tutorial()
    print("\nPart 1 complete (3/6 documents generated).")
