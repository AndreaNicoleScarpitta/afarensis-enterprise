"""
Afarensis Enterprise — Analysis Data Reviewer's Guide (ADRG) Generator

Generates the ADRG document per FDA and PhUSE recommendations.
The ADRG provides FDA reviewers with instructions for recreating
the analysis environment and reproducing all statistical outputs.

Critical for submissions using R or Python (non-SAS).
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import io
from datetime import datetime
from typing import Dict, Any, Optional
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select
import logging

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Default XY-301 study data (mirrors document_generator.py defaults)
# ---------------------------------------------------------------------------
_XY301_DEFAULTS = {
    "protocol": "XY-301",
    "title": "External Control Arm Study for Rare CNS Disorder (Pediatric)",
    "indication": "Rare CNS Disorder (Pediatric)",
    "sponsor": "Afarensis Therapeutics, Inc.",
    "phase": "Phase 3",
    "primary_endpoint": "All-cause hospitalization (time-to-first event)",
    "study_design": "Single-arm trial with external control arm",
    "sample_size_trial": 112,
    "sample_size_eca": 489,
}


class ADRGGenerator:
    """Generates Analysis Data Reviewer's Guide (ADRG) documents."""

    NAVY = RGBColor(0x1B, 0x2A, 0x4A)

    # ------------------------------------------------------------------
    # Internal helpers (match document_generator.py conventions)
    # ------------------------------------------------------------------

    @staticmethod
    def _now_str() -> str:
        return datetime.utcnow().strftime("%d %B %Y")

    def _val(self, project_data: Optional[dict], key: str, fallback=None):
        """Pull a value from project_data (nested under study_definition or
        processing_config) with XY-301 defaults as final fallback."""
        if project_data:
            # direct key
            if key in project_data:
                return project_data[key]
            # nested under study_definition
            sd = project_data.get("study_definition", {})
            if key in sd:
                return sd[key]
            # nested under processing_config
            pc = project_data.get("processing_config", {})
            if key in pc:
                return pc[key]
        return fallback if fallback is not None else _XY301_DEFAULTS.get(key, "")

    # ------------------------------------------------------------------
    # Public entry point
    # ------------------------------------------------------------------

    async def generate_adrg_docx(
        self,
        db: AsyncSession,
        project_id: str,
        project_data: dict = None,
    ) -> bytes:
        """
        Generate the ADRG as a DOCX document.

        Follows PhUSE Analysis Data Reviewer's Guide v1.1 structure.
        """
        from io import BytesIO

        pd = project_data or {}
        protocol = self._val(pd, "protocol", _XY301_DEFAULTS["protocol"])
        title = self._val(pd, "title", _XY301_DEFAULTS["title"])
        sponsor = self._val(pd, "sponsor", _XY301_DEFAULTS["sponsor"])
        indication = self._val(pd, "indication", _XY301_DEFAULTS["indication"])
        phase = self._val(pd, "phase", _XY301_DEFAULTS["phase"])
        primary_endpoint = self._val(pd, "primary_endpoint", _XY301_DEFAULTS["primary_endpoint"])
        study_design = self._val(pd, "study_design", _XY301_DEFAULTS["study_design"])
        n_trial = self._val(pd, "sample_size_trial", _XY301_DEFAULTS["sample_size_trial"])
        n_eca = self._val(pd, "sample_size_eca", _XY301_DEFAULTS["sample_size_eca"])
        now = self._now_str()

        doc = Document()
        style = doc.styles["Normal"]
        font = style.font
        font.name = "Times New Roman"
        font.size = Pt(11)

        navy = self.NAVY

        # ---- reusable helpers (same pattern as SAR/SAP) ----
        def add_heading_styled(text, level=1):
            h = doc.add_heading(text, level=level)
            for run in h.runs:
                run.font.color.rgb = navy
            return h

        def add_table_from_rows(headers, rows):
            tbl = doc.add_table(rows=1, cols=len(headers))
            tbl.style = "Light Grid Accent 1"
            tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
            hdr_cells = tbl.rows[0].cells
            for i, h in enumerate(headers):
                hdr_cells[i].text = h
                for par in hdr_cells[i].paragraphs:
                    for run in par.runs:
                        run.bold = True
                        run.font.size = Pt(9)
            for row_data in rows:
                row_cells = tbl.add_row().cells
                for i, val in enumerate(row_data):
                    row_cells[i].text = str(val) if val is not None else "\u2014"
                    for par in row_cells[i].paragraphs:
                        for run in par.runs:
                            run.font.size = Pt(9)
            return tbl

        # ==============================================================
        # 1. Title Page
        # ==============================================================
        doc.add_paragraph("")
        doc.add_paragraph("")
        title_para = doc.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title_para.add_run("ANALYSIS DATA REVIEWER\u2019S GUIDE")
        run.bold = True
        run.font.size = Pt(26)
        run.font.color.rgb = navy

        sub = doc.add_paragraph()
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sr = sub.add_run(f"Protocol {protocol} \u2014 {title}")
        sr.font.size = Pt(14)
        sr.font.color.rgb = RGBColor(0x4A, 0x4A, 0x4A)

        doc.add_paragraph("")
        for label, val in [
            ("Study", protocol),
            ("Sponsor", sponsor),
            ("Version", "1.0"),
            ("Date", now),
            ("Prepared By", "Biostatistics / Programming Department"),
        ]:
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r1 = para.add_run(f"{label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = para.add_run(str(val))
            r2.font.size = Pt(11)

        doc.add_paragraph("")
        conf = doc.add_paragraph()
        conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cr = conf.add_run("CONFIDENTIAL \u2014 FOR REGULATORY USE ONLY")
        cr.bold = True
        cr.font.size = Pt(10)
        cr.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

        doc.add_page_break()

        # ==============================================================
        # 2. Table of Contents (placeholder)
        # ==============================================================
        add_heading_styled("Table of Contents", level=1)
        doc.add_paragraph(
            "[Insert Table of Contents \u2014 update field after final document review]"
        )
        doc.add_page_break()

        # ==============================================================
        # 3. Introduction
        # ==============================================================
        add_heading_styled("1. Introduction", level=1)

        add_heading_styled("1.1 Purpose", level=2)
        doc.add_paragraph(
            "This Analysis Data Reviewer\u2019s Guide (ADRG) provides FDA reviewers with "
            "the information needed to understand and reproduce the analysis datasets "
            f"and statistical outputs for Protocol {protocol}. It describes the "
            "computational environment, analysis programs, and data flow from source "
            "data through final Tables, Figures, and Listings (TFLs)."
        )

        add_heading_styled("1.2 Scope and Intended Audience", level=2)
        doc.add_paragraph(
            "This guide is intended for FDA statistical reviewers and programmers "
            "who need to recreate the analysis environment and verify the submitted "
            "results. It covers the complete chain from raw SDTM datasets through "
            "ADaM analysis datasets to the final statistical outputs."
        )

        add_heading_styled("1.3 Related Documents", level=2)
        for ref_doc in [
            f"Statistical Analysis Plan (SAP), Protocol {protocol}",
            "Define-XML v2.0 \u2014 Analysis Dataset Metadata",
            "Annotated Case Report Forms (aCRF)",
            "Clinical Study Report (CSR)",
            "SDTM Reviewer\u2019s Guide (SDRG)",
        ]:
            doc.add_paragraph(ref_doc, style="List Bullet")

        doc.add_page_break()

        # ==============================================================
        # 4. Study Overview
        # ==============================================================
        add_heading_styled("2. Study Overview", level=1)

        add_heading_styled("2.1 Study Information", level=2)
        add_table_from_rows(
            ["Attribute", "Value"],
            [
                ["Study Title", title],
                ["Protocol Number", protocol],
                ["Phase", phase],
                ["Indication", indication],
                ["Sponsor", sponsor],
                ["Study Design", study_design],
                ["Primary Endpoint", primary_endpoint],
            ],
        )

        add_heading_styled("2.2 Sample Size Summary", level=2)
        add_table_from_rows(
            ["Cohort", "N"],
            [
                [f"Trial Arm ({protocol})", n_trial],
                ["External Control Arm", n_eca],
                ["Total", int(n_trial) + int(n_eca)],
            ],
        )

        doc.add_page_break()

        # ==============================================================
        # 5. Analysis Datasets
        # ==============================================================
        add_heading_styled("3. Analysis Datasets", level=1)

        doc.add_paragraph(
            "The following ADaM datasets are used in the analysis. All datasets "
            "conform to CDISC ADaM Implementation Guide v1.3 and are derived from "
            "the corresponding SDTM domains."
        )

        add_heading_styled("3.1 Dataset Inventory", level=2)
        add_table_from_rows(
            ["Dataset", "Description", "Structure", "Key Variables", "Records"],
            [
                ["ADSL", "Subject-Level Analysis Dataset",
                 "One record per subject",
                 "USUBJID, TRT01P, AGE, SEX, RACE, ITTFL, SAFFL",
                 str(int(n_trial) + int(n_eca))],
                ["ADAE", "Analysis Dataset for Adverse Events",
                 "One record per AE per subject",
                 "USUBJID, AEDECOD, AESEV, AESER, TRTEMFL",
                 "\u2014"],
                ["ADTTE", "Analysis Dataset for Time-to-Event",
                 "One record per parameter per subject",
                 "USUBJID, PARAMCD, AVAL, CNSR, STARTDT, ADT",
                 "\u2014"],
                ["ADLB", "Analysis Dataset for Lab Results",
                 "One record per lab test per visit per subject",
                 "USUBJID, PARAMCD, AVAL, BASE, CHG, AVISITN",
                 "\u2014"],
            ],
        )

        add_heading_styled("3.2 SDTM to ADaM Traceability", level=2)
        doc.add_paragraph(
            "All ADaM datasets maintain full traceability to their SDTM source "
            "domains. The primary traceability paths are:"
        )
        for trace in [
            "DM + EX + DS \u2192 ADSL (subject-level variables, disposition, exposure)",
            "AE + ADSL \u2192 ADAE (adverse event analysis flags, treatment-emergent classification)",
            "AE + EX + ADSL \u2192 ADTTE (time-to-event derivations for primary endpoint)",
            "LB + ADSL \u2192 ADLB (laboratory shift analysis, baseline derivations)",
        ]:
            doc.add_paragraph(trace, style="List Bullet")

        doc.add_page_break()

        # ==============================================================
        # 6. Analysis Software
        # ==============================================================
        add_heading_styled("4. Analysis Software", level=1)

        doc.add_paragraph(
            "All analyses were conducted using open-source software. The table "
            "below lists all software and package versions used in the analysis."
        )

        add_heading_styled("4.1 Software Versions", level=2)
        add_table_from_rows(
            ["Software / Package", "Version", "Purpose"],
            [
                ["Python", "3.11", "Primary analysis engine"],
                ["NumPy", "1.26.4", "Numerical computation"],
                ["SciPy", "1.14.1", "Statistical tests (log-rank, chi-square, t-tests)"],
                ["statsmodels", "0.14.4", "Mixed models, regression, survival analysis"],
                ["pandas", "2.2.3", "Data manipulation and transformation"],
                ["matplotlib", "3.10.0", "Figure and plot generation"],
                ["python-docx", "1.1.2", "Regulatory document generation (SAR, SAP, CSR)"],
                ["scikit-learn", "1.5.2", "Propensity score estimation (logistic regression)"],
                ["lifelines", "0.29.0", "Kaplan-Meier estimation, Cox PH models"],
            ],
        )

        doc.add_paragraph("")
        note_para = doc.add_paragraph()
        r = note_para.add_run("Note: ")
        r.bold = True
        r.font.size = Pt(10)
        note_para.add_run(
            "All packages are available via pip install. A complete frozen "
            "environment specification is provided in requirements.txt."
        ).font.size = Pt(10)

        doc.add_page_break()

        # ==============================================================
        # 7. Environment Recreation Instructions
        # ==============================================================
        add_heading_styled("5. Environment Recreation Instructions", level=1)

        doc.add_paragraph(
            "The following step-by-step instructions allow a reviewer to recreate "
            "the computational environment and reproduce all analysis outputs."
        )

        add_heading_styled("5.1 Standard Installation (pip)", level=2)
        steps_pip = [
            ("Step 1", "Install Python 3.11 from https://www.python.org/downloads/"),
            ("Step 2", "Create a virtual environment:\n"
             "    python -m venv afarensis_env\n"
             "    source afarensis_env/bin/activate  (Linux/Mac)\n"
             "    afarensis_env\\Scripts\\activate  (Windows)"),
            ("Step 3", "Install dependencies:\n"
             "    pip install -r requirements.txt"),
            ("Step 4", "Run the analysis:\n"
             "    python -m app.services.statistical_models"),
            ("Step 5", "Verify outputs match archived results (see Section 7)."),
        ]
        for step_label, step_text in steps_pip:
            para = doc.add_paragraph()
            r1 = para.add_run(f"{step_label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            para.add_run(step_text).font.size = Pt(11)

        add_heading_styled("5.2 Docker Installation (Recommended)", level=2)
        doc.add_paragraph(
            "A pre-built Docker image is available that contains the complete "
            "analysis environment with all dependencies locked to exact versions."
        )
        steps_docker = [
            ("Step 1", f"Pull the Docker image:\n"
             f"    docker pull afarensis/xy301:v1.0"),
            ("Step 2", "Run the container:\n"
             f"    docker run -v /path/to/data:/data afarensis/xy301:v1.0"),
            ("Step 3", "Outputs are written to /data/output/ inside the container."),
        ]
        for step_label, step_text in steps_docker:
            para = doc.add_paragraph()
            r1 = para.add_run(f"{step_label}: ")
            r1.bold = True
            r1.font.size = Pt(11)
            para.add_run(step_text).font.size = Pt(11)

        add_heading_styled("5.3 Conda Environment", level=2)
        doc.add_paragraph(
            "Alternatively, a conda environment file is provided:\n"
            "    conda env create -f environment.yml\n"
            "    conda activate afarensis_xy301"
        )

        doc.add_page_break()

        # ==============================================================
        # 8. Analysis Program Descriptions
        # ==============================================================
        add_heading_styled("6. Analysis Program Descriptions", level=1)

        doc.add_paragraph(
            "The following table describes each analysis program, its purpose, "
            "input datasets, and the output TFLs it produces."
        )

        add_table_from_rows(
            ["Program File", "Description", "Input Datasets", "Output TFLs"],
            [
                ["statistical_models.py",
                 "Primary and sensitivity analyses (IPTW Cox PH, PS matching, "
                 "AIPW, overlap weighting, tipping-point analysis)",
                 "ADSL, ADTTE",
                 "T-14.2.1, T-14.2.2, T-14.2.3"],
                ["tfl_generator.py",
                 "Table and figure generation (Kaplan-Meier curves, forest plots, "
                 "Love plots, demographics tables)",
                 "ADSL, ADAE, ADTTE, ADLB",
                 "T-14.1.1, T-14.1.2, T-14.1.3, T-14.3.1, T-14.3.2, "
                 "F-14.1.1, F-14.1.2, F-14.1.3, F-14.1.4"],
                ["document_generator.py",
                 "Regulatory document assembly (SAR, SAP, CSR as DOCX and HTML)",
                 "Analysis results JSON",
                 "SAR document, SAP document, CSR sections"],
                ["adrg_generator.py",
                 "Analysis Data Reviewer\u2019s Guide generation (this document)",
                 "Project metadata",
                 "ADRG document"],
                ["csr_generator.py",
                 "Clinical Study Report section generation (ICH E3 Sections 11, 12, 16)",
                 "ADSL, ADAE, ADTTE, analysis results",
                 "CSR Synopsis, Section 11, Section 12, Appendix 16"],
            ],
        )

        doc.add_page_break()

        # ==============================================================
        # 9. Analysis Results Verification
        # ==============================================================
        add_heading_styled("7. Analysis Results Verification", level=1)

        add_heading_styled("7.1 Verification Approach", level=2)
        doc.add_paragraph(
            "All analysis outputs are accompanied by SHA-256 checksums. Reviewers "
            "should regenerate outputs using the instructions in Section 5 and "
            "compare checksums to verify exact reproducibility."
        )

        add_heading_styled("7.2 Expected Output Files and Checksums", level=2)
        add_table_from_rows(
            ["Output File", "Description", "Checksum (SHA-256)"],
            [
                ["results/primary_analysis.json", "Primary IPTW Cox PH results", "[computed at submission]"],
                ["results/sensitivity_analyses.json", "All sensitivity analysis results", "[computed at submission]"],
                ["results/subgroup_analyses.json", "Subgroup analysis results", "[computed at submission]"],
                ["output/tables/T-14.2.1.docx", "Primary analysis table", "[computed at submission]"],
                ["output/figures/F-14.1.1.png", "Kaplan-Meier plot", "[computed at submission]"],
                ["output/figures/F-14.1.2.png", "Forest plot", "[computed at submission]"],
            ],
        )

        add_heading_styled("7.3 Numerical Comparison Tolerances", level=2)
        doc.add_paragraph(
            "For numerical results, the following tolerances are applied when "
            "comparing reproduced outputs to archived results:"
        )
        add_table_from_rows(
            ["Result Type", "Tolerance", "Rationale"],
            [
                ["Point estimates (HR, OR)", "1e-10", "Deterministic computation; exact match expected"],
                ["P-values", "1e-10", "Deterministic computation; exact match expected"],
                ["Confidence interval bounds", "1e-10", "Deterministic computation"],
                ["Bootstrap-based estimates", "1e-4", "Random seed-dependent; seed is fixed but platform "
                 "differences may introduce minor variation"],
                ["Figure pixel values", "N/A", "Visual comparison; minor rendering differences acceptable"],
            ],
        )

        doc.add_page_break()

        # ==============================================================
        # 10. Data Flow Diagrams
        # ==============================================================
        add_heading_styled("8. Data Flow Diagrams", level=1)

        doc.add_paragraph(
            "The complete data pipeline from raw source data to final regulatory "
            "outputs follows the stages described below."
        )

        add_heading_styled("8.1 Pipeline Overview", level=2)
        stages = [
            ("Stage 1: Raw Data Collection",
             "Clinical trial CRF data and real-world data (claims + EHR) are "
             "collected and quality-checked.",
             "Input: Raw CRF exports, claims extracts, EHR pulls\n"
             "Output: Cleaned raw datasets"),
            ("Stage 2: SDTM Conversion",
             "Raw data are mapped to CDISC SDTM domains (DM, AE, EX, DS, LB).",
             "Input: Cleaned raw datasets, SDTM mapping specifications\n"
             "Output: SDTM datasets (XPT format)"),
            ("Stage 3: ADaM Derivation",
             "SDTM domains are transformed into ADaM analysis datasets.",
             "Input: SDTM datasets, ADaM specifications\n"
             "Output: ADSL, ADAE, ADTTE, ADLB (XPT and SAS7BDAT)"),
            ("Stage 4: Statistical Analysis",
             "Primary, sensitivity, and subgroup analyses are executed.",
             "Input: ADaM datasets\n"
             "Output: Analysis results (JSON), model objects"),
            ("Stage 5: TFL Generation",
             "Tables, figures, and listings are produced from analysis results.",
             "Input: Analysis results, ADaM datasets\n"
             "Output: TFLs (DOCX, PNG, PDF)"),
            ("Stage 6: Document Assembly",
             "Regulatory documents (SAR, SAP, CSR) are assembled.",
             "Input: TFLs, analysis results, project metadata\n"
             "Output: SAR, SAP, CSR, ADRG (DOCX and HTML)"),
        ]
        for stage_title, description, io_spec in stages:
            para = doc.add_paragraph()
            r1 = para.add_run(stage_title)
            r1.bold = True
            r1.font.size = Pt(11)
            doc.add_paragraph(description)
            io_para = doc.add_paragraph()
            io_para.add_run(io_spec).font.size = Pt(10)
            doc.add_paragraph("")

        doc.add_page_break()

        # ==============================================================
        # 11. Special Considerations
        # ==============================================================
        add_heading_styled("9. Special Considerations", level=1)

        add_heading_styled("9.1 Random Seed Usage", level=2)
        doc.add_paragraph(
            "All stochastic procedures use explicitly set random seeds to ensure "
            "reproducibility. The global random seed is set to 42 at the beginning "
            "of each analysis script. Procedure-specific seeds are derived "
            "deterministically from the global seed."
        )
        doc.add_paragraph(
            "Procedures using random seeds include: bootstrap confidence intervals, "
            "multiple imputation, propensity score matching (caliper-based neighbor "
            "selection tie-breaking), and cross-validation for GBM tuning."
        )

        add_heading_styled("9.2 Platform-Dependent Numerical Differences", level=2)
        doc.add_paragraph(
            "Minor numerical differences (typically < 1e-14) may arise due to "
            "differences in floating-point arithmetic across platforms (e.g., "
            "Intel vs. ARM, different BLAS/LAPACK implementations). These differences "
            "are within machine epsilon and do not affect scientific conclusions."
        )
        doc.add_paragraph(
            "The reference platform for this submission is: Linux x86_64, "
            "Intel MKL 2024.0, Python 3.11.7, compiled with GCC 12.3."
        )

        add_heading_styled("9.3 Missing Data Handling", level=2)
        doc.add_paragraph(
            "The primary analysis assumes a Missing at Random (MAR) mechanism. "
            "Missing covariate values are handled via multiple imputation (m = 20 "
            "datasets, Rubin\u2019s rules for pooling). Missing outcome data are "
            "addressed through the censoring mechanism in the Cox PH model."
        )
        doc.add_paragraph(
            "Sensitivity analyses assess departure from MAR via tipping-point "
            "analysis (delta adjustment applied to external control arm outcomes)."
        )

        doc.add_page_break()

        # ==============================================================
        # 12. References
        # ==============================================================
        add_heading_styled("10. References", level=1)
        for ref in [
            "PhUSE Analysis Data Reviewer\u2019s Guide (ADRG) v1.1, 2022.",
            "FDA Study Data Technical Conformance Guide v4.6, 2023.",
            "CDISC Analysis Data Model (ADaM) Implementation Guide v1.3, 2021.",
            "CDISC Study Data Tabulation Model (SDTM) Implementation Guide v3.4, 2021.",
            "ICH E3: Structure and Content of Clinical Study Reports, 1995.",
            "ICH E9: Statistical Principles for Clinical Trials, 1998.",
            "ICH E9(R1): Addendum on Estimands and Sensitivity Analysis in "
            "Clinical Trials, 2019.",
            "FDA Guidance: Considerations for the Design and Conduct of Externally "
            "Controlled Trials for Drug and Biological Products, 2023.",
            "FDA Guidance: Providing Regulatory Submissions in Electronic Format \u2014 "
            "Standardized Study Data, 2023.",
        ]:
            doc.add_paragraph(ref, style="List Number")

        # ---- Footer ----
        doc.add_paragraph("")
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = footer_para.add_run(
            f"Generated by Afarensis Enterprise v2.1 \u2014 {now} \u2014 "
            f"Protocol {protocol} \u2014 CONFIDENTIAL"
        )
        fr.font.size = Pt(8)
        fr.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

        # ---- serialize ----
        buf = BytesIO()
        doc.save(buf)
        return buf.getvalue()
